from __future__ import annotations

from pathlib import Path
import json
import zipfile
from html import escape

import numpy as np
import pandas as pd
from scipy.stats import wilcoxon

import matplotlib

matplotlib.use("Agg")
import matplotlib.pyplot as plt


import sys
from pathlib import Path

_SI_ROOT = Path(__file__).resolve().parents[1]
if str(_SI_ROOT) not in sys.path:
    sys.path.insert(0, str(_SI_ROOT))
from _paths import (  # noqa: E402
    BENCHMARK,
    CONFIG_YAML,
    DATASET_META_JSON,
    FINAL,
    INTERVALS_CSV,
    OUT,
    PROCESSED_CSV,
    RESULTS_CSV,
    ROOT,
    RUN,
    S8_MANIFEST,
    SIMILARITY_CSV,
    SUMMARY_CSV,
    XGB_SHAP_CSV,
    XGB_SHAP_FIG,
)


RESULTS_CSV = RUN / "metrics" / "results.csv"
SUMMARY_CSV = RUN / "metrics" / "summary_frac1_alpha.csv"
DATASET_META_JSON = RUN / "data" / "dataset_meta.json"
RUN_META_JSON = RUN / "run_meta.json"

MAIN_FRAC = 1.0
MAIN_ALPHA = 0.10
BOOTSTRAP_N = 20000
RNG_SEED = 20260508


METRIC_INFO = {
    "RMSE": {"label": "RMSE (K)", "higher_better": False},
    "MAE": {"label": "MAE (K)", "higher_better": False},
    "R2": {"label": "R2", "higher_better": True},
    "width": {"label": "Conformal mean interval width (K)", "higher_better": False},
}


def p_format(p: float) -> str:
    if pd.isna(p):
        return "NA"
    if p < 0.001:
        return "<0.001"
    return f"{p:.3f}"


def ci_format(low: float, high: float, decimals: int = 2) -> str:
    return f"[{low:.{decimals}f}, {high:.{decimals}f}]"


def bootstrap_mean_ci(diff: np.ndarray, n_boot: int, seed: int) -> tuple[float, float]:
    rng = np.random.default_rng(seed)
    diff = np.asarray(diff, dtype=float)
    boot = rng.choice(diff, size=(n_boot, diff.size), replace=True).mean(axis=1)
    return tuple(np.percentile(boot, [2.5, 97.5]))


def wilcoxon_test(diff: np.ndarray) -> tuple[float, float]:
    diff = np.asarray(diff, dtype=float)
    if np.allclose(diff, 0):
        return 0.0, 1.0
    res = wilcoxon(diff, zero_method="wilcox", alternative="two-sided", method="auto")
    return float(res.statistic), float(res.pvalue)


def matched_pairs(results: pd.DataFrame, frac: float, alpha: float) -> pd.DataFrame:
    df = results[(results["frac"] == frac) & (results["alpha"] == alpha)].copy()
    df["cutoff_key"] = df["cutoff"].fillna("none").astype(str)
    keys = ["regime", "base_regime", "cutoff_key", "seed", "fold", "frac", "alpha"]
    value_cols = ["RMSE", "MAE", "R2", "width", "cov", "nominal_coverage"]
    wide = df.pivot_table(index=keys, columns="model", values=value_cols, aggfunc="first")
    wide.columns = [f"{metric}_{model}" for metric, model in wide.columns]
    wide = wide.reset_index()
    wide["cutoff"] = pd.to_numeric(wide["cutoff_key"].replace("none", np.nan), errors="coerce")

    required = []
    for metric in METRIC_INFO:
        required.extend([f"{metric}_svr", f"{metric}_xgb"])
    missing = [col for col in required if col not in wide.columns]
    if missing:
        raise ValueError(f"Missing matched model columns: {missing}")
    wide = wide.dropna(subset=required)
    return wide


def paired_table(wide: pd.DataFrame) -> tuple[pd.DataFrame, pd.DataFrame]:
    rows = []
    pair_rows = []
    groups = [("Overall", wide)]
    groups.extend((regime, g) for regime, g in wide.groupby("regime", sort=True))

    for group_name, group in groups:
        for metric, info in METRIC_INFO.items():
            svr = group[f"{metric}_svr"].to_numpy(dtype=float)
            xgb = group[f"{metric}_xgb"].to_numpy(dtype=float)
            diff = svr - xgb
            ci_low, ci_high = bootstrap_mean_ci(
                diff, BOOTSTRAP_N, RNG_SEED + len(rows) * 17
            )
            stat, p = wilcoxon_test(diff)
            mean_diff = float(np.mean(diff))
            if info["higher_better"]:
                better = "SVR" if mean_diff > 0 else "XGBoost"
            else:
                better = "SVR" if mean_diff < 0 else "XGBoost"
            rows.append(
                {
                    "comparison_set": group_name,
                    "metric": metric,
                    "metric_label": info["label"],
                    "n_pairs": int(group.shape[0]),
                    "SVR_mean": float(np.mean(svr)),
                    "XGBoost_mean": float(np.mean(xgb)),
                    "paired_mean_difference_SVR_minus_XGBoost": mean_diff,
                    "bootstrap_95CI_low": float(ci_low),
                    "bootstrap_95CI_high": float(ci_high),
                    "wilcoxon_W": stat,
                    "wilcoxon_p": p,
                    "lower_is_better": not info["higher_better"],
                    "favored_by_mean_difference": better,
                }
            )

        for _, r in group.iterrows():
            base = {
                "comparison_set": group_name,
                "regime": r["regime"],
                "base_regime": r["base_regime"],
                "cutoff": r["cutoff"],
                "seed": int(r["seed"]),
                "fold": int(r["fold"]),
                "frac": float(r["frac"]),
                "alpha": float(r["alpha"]),
            }
            for metric in METRIC_INFO:
                pair_rows.append(
                    {
                        **base,
                        "metric": metric,
                        "SVR": float(r[f"{metric}_svr"]),
                        "XGBoost": float(r[f"{metric}_xgb"]),
                        "paired_difference_SVR_minus_XGBoost": float(
                            r[f"{metric}_svr"] - r[f"{metric}_xgb"]
                        ),
                    }
                )

    return pd.DataFrame(rows), pd.DataFrame(pair_rows)


def alpha_width_sensitivity(results: pd.DataFrame) -> pd.DataFrame:
    rows = []
    for alpha, alpha_df in results[results["frac"] == MAIN_FRAC].groupby("alpha"):
        wide = matched_pairs(results, MAIN_FRAC, float(alpha))
        for group_name, group in [("Overall", wide)] + list(wide.groupby("regime", sort=True)):
            diff = group["width_svr"].to_numpy(float) - group["width_xgb"].to_numpy(float)
            ci_low, ci_high = bootstrap_mean_ci(
                diff, BOOTSTRAP_N, RNG_SEED + int(float(alpha) * 1000)
            )
            stat, p = wilcoxon_test(diff)
            rows.append(
                {
                    "comparison_set": group_name,
                    "alpha": float(alpha),
                    "nominal_coverage": float(1.0 - alpha),
                    "n_pairs": int(group.shape[0]),
                    "SVR_width_mean": float(group["width_svr"].mean()),
                    "XGBoost_width_mean": float(group["width_xgb"].mean()),
                    "paired_mean_difference_SVR_minus_XGBoost": float(diff.mean()),
                    "bootstrap_95CI_low": float(ci_low),
                    "bootstrap_95CI_high": float(ci_high),
                    "wilcoxon_W": stat,
                    "wilcoxon_p": p,
                }
            )
    return pd.DataFrame(rows)


def manuscript_view(table: pd.DataFrame) -> pd.DataFrame:
    out = table.copy()
    out["SVR mean"] = out["SVR_mean"].map(lambda x: f"{x:.2f}")
    out["XGBoost mean"] = out["XGBoost_mean"].map(lambda x: f"{x:.2f}")
    out["Mean difference (SVR-XGBoost)"] = out[
        "paired_mean_difference_SVR_minus_XGBoost"
    ].map(lambda x: f"{x:.2f}")
    out["Bootstrap 95% CI"] = [
        ci_format(lo, hi, decimals=2)
        for lo, hi in zip(out["bootstrap_95CI_low"], out["bootstrap_95CI_high"])
    ]
    out["Wilcoxon W"] = out["wilcoxon_W"].map(lambda x: f"{x:.1f}")
    out["Wilcoxon p"] = out["wilcoxon_p"].map(p_format)
    out["Direction"] = out["favored_by_mean_difference"]
    return out[
        [
            "comparison_set",
            "metric_label",
            "n_pairs",
            "SVR mean",
            "XGBoost mean",
            "Mean difference (SVR-XGBoost)",
            "Bootstrap 95% CI",
            "Wilcoxon W",
            "Wilcoxon p",
            "Direction",
        ]
    ].rename(
        columns={
            "comparison_set": "Comparison set",
            "metric_label": "Metric",
            "n_pairs": "Matched pairs",
        }
    )


def dataframe_to_markdown(df: pd.DataFrame) -> str:
    headers = [str(col) for col in df.columns]
    rows = [[str(value) for value in row] for row in df.to_numpy()]
    widths = [
        max(len(headers[i]), *(len(row[i]) for row in rows)) if rows else len(headers[i])
        for i in range(len(headers))
    ]
    header_line = "| " + " | ".join(headers[i].ljust(widths[i]) for i in range(len(headers))) + " |"
    separator = "| " + " | ".join("-" * widths[i] for i in range(len(headers))) + " |"
    body = [
        "| " + " | ".join(row[i].ljust(widths[i]) for i in range(len(headers))) + " |"
        for row in rows
    ]
    return "\n".join([header_line, separator, *body]) + "\n"


def latex_escape(value: object) -> str:
    text = str(value)
    replacements = {
        "\\": r"\textbackslash{}",
        "&": r"\&",
        "%": r"\%",
        "$": r"\$",
        "#": r"\#",
        "_": r"\_",
        "{": r"\{",
        "}": r"\}",
        "~": r"\textasciitilde{}",
        "^": r"\textasciicircum{}",
    }
    for old, new in replacements.items():
        text = text.replace(old, new)
    return text


def dataframe_to_latex(df: pd.DataFrame) -> str:
    cols = list(df.columns)
    lines = [
        r"\begin{tabular}{" + "l" * len(cols) + "}",
        r"\hline",
        " & ".join(latex_escape(col) for col in cols) + r" \\",
        r"\hline",
    ]
    for _, row in df.iterrows():
        lines.append(" & ".join(latex_escape(row[col]) for col in cols) + r" \\")
    lines.extend([r"\hline", r"\end{tabular}", ""])
    return "\n".join(lines)


def excel_col_name(index: int) -> str:
    name = ""
    while index:
        index, remainder = divmod(index - 1, 26)
        name = chr(65 + remainder) + name
    return name


def worksheet_xml(df: pd.DataFrame) -> str:
    rows = [list(df.columns)] + df.astype(str).values.tolist()
    xml_rows = []
    for r_idx, row in enumerate(rows, start=1):
        cells = []
        for c_idx, value in enumerate(row, start=1):
            ref = f"{excel_col_name(c_idx)}{r_idx}"
            cells.append(
                f'<c r="{ref}" t="inlineStr"><is><t>{escape(str(value))}</t></is></c>'
            )
        xml_rows.append(f'<row r="{r_idx}">{"".join(cells)}</row>')
    return (
        '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
        '<worksheet xmlns="http://schemas.openxmlformats.org/spreadsheetml/2006/main">'
        f'<sheetData>{"".join(xml_rows)}</sheetData>'
        "</worksheet>"
    )


def write_xlsx(path: Path, sheets: dict[str, pd.DataFrame]) -> None:
    sheet_names = list(sheets.keys())
    with zipfile.ZipFile(path, "w", zipfile.ZIP_DEFLATED) as zf:
        zf.writestr(
            "[Content_Types].xml",
            '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
            '<Types xmlns="http://schemas.openxmlformats.org/package/2006/content-types">'
            '<Default Extension="rels" ContentType="application/vnd.openxmlformats-package.relationships+xml"/>'
            '<Default Extension="xml" ContentType="application/xml"/>'
            '<Override PartName="/xl/workbook.xml" ContentType="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet.main+xml"/>'
            + "".join(
                f'<Override PartName="/xl/worksheets/sheet{i}.xml" '
                'ContentType="application/vnd.openxmlformats-officedocument.spreadsheetml.worksheet+xml"/>'
                for i in range(1, len(sheet_names) + 1)
            )
            + "</Types>",
        )
        zf.writestr(
            "_rels/.rels",
            '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
            '<Relationships xmlns="http://schemas.openxmlformats.org/package/2006/relationships">'
            '<Relationship Id="rId1" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/officeDocument" Target="xl/workbook.xml"/>'
            "</Relationships>",
        )
        zf.writestr(
            "xl/workbook.xml",
            '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
            '<workbook xmlns="http://schemas.openxmlformats.org/spreadsheetml/2006/main" '
            'xmlns:r="http://schemas.openxmlformats.org/officeDocument/2006/relationships">'
            "<sheets>"
            + "".join(
                f'<sheet name="{escape(name[:31])}" sheetId="{i}" r:id="rId{i}"/>'
                for i, name in enumerate(sheet_names, start=1)
            )
            + "</sheets></workbook>",
        )
        zf.writestr(
            "xl/_rels/workbook.xml.rels",
            '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
            '<Relationships xmlns="http://schemas.openxmlformats.org/package/2006/relationships">'
            + "".join(
                f'<Relationship Id="rId{i}" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/worksheet" '
                f'Target="worksheets/sheet{i}.xml"/>'
                for i in range(1, len(sheet_names) + 1)
            )
            + "</Relationships>",
        )
        for i, name in enumerate(sheet_names, start=1):
            zf.writestr(f"xl/worksheets/sheet{i}.xml", worksheet_xml(sheets[name]))


def docx_text(text: str, style: str | None = None) -> str:
    style_xml = f'<w:pPr><w:pStyle w:val="{style}"/></w:pPr>' if style else ""
    return f"<w:p>{style_xml}<w:r><w:t>{escape(text)}</w:t></w:r></w:p>"


def write_minimal_docx_table(path: Path, manuscript: pd.DataFrame) -> None:
    rows = [list(manuscript.columns)] + manuscript.astype(str).values.tolist()
    table_rows = []
    for row in rows:
        cells = "".join(
            "<w:tc><w:p><w:r><w:t>"
            + escape(str(value))
            + "</w:t></w:r></w:p></w:tc>"
            for value in row
        )
        table_rows.append(f"<w:tr>{cells}</w:tr>")
    document_xml = (
        '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
        '<w:document xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
        "<w:body>"
        + docx_text(
            "Table S7. Paired statistical comparison of SVR and XGBoost across matched seed/fold evaluations",
            "Heading1",
        )
        + docx_text(
            "Values are based on full-training evaluations (learning fraction = 1.0) and "
            "90% nominal split-conformal intervals (alpha = 0.10). Paired differences are "
            "SVR minus XGBoost. For RMSE, MAE, and interval width, negative differences "
            "favor SVR; for R2, positive differences favor SVR."
        )
        + f"<w:tbl>{''.join(table_rows)}</w:tbl>"
        + '<w:sectPr><w:pgSz w:w="15840" w:h="12240" w:orient="landscape"/></w:sectPr>'
        + "</w:body></w:document>"
    )
    with zipfile.ZipFile(path, "w", zipfile.ZIP_DEFLATED) as zf:
        zf.writestr(
            "[Content_Types].xml",
            '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
            '<Types xmlns="http://schemas.openxmlformats.org/package/2006/content-types">'
            '<Default Extension="rels" ContentType="application/vnd.openxmlformats-package.relationships+xml"/>'
            '<Default Extension="xml" ContentType="application/xml"/>'
            '<Override PartName="/word/document.xml" ContentType="application/vnd.openxmlformats-officedocument.wordprocessingml.document.main+xml"/>'
            "</Types>",
        )
        zf.writestr(
            "_rels/.rels",
            '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
            '<Relationships xmlns="http://schemas.openxmlformats.org/package/2006/relationships">'
            '<Relationship Id="rId1" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/officeDocument" Target="word/document.xml"/>'
            "</Relationships>",
        )
        zf.writestr("word/document.xml", document_xml)


def write_docx_table(path: Path, manuscript: pd.DataFrame) -> bool:
    try:
        from docx import Document
    except Exception:
        write_minimal_docx_table(path, manuscript)
        return True

    doc = Document()
    doc.add_heading(
        "Table S7. Paired statistical comparison of SVR and XGBoost across matched seed/fold evaluations",
        level=1,
    )
    doc.add_paragraph(
        "Values are based on full-training evaluations (learning fraction = 1.0) "
        "and 90% nominal split-conformal intervals (alpha = 0.10). "
        "Paired differences are SVR minus XGBoost. For RMSE, MAE, and interval "
        "width, negative differences favor SVR; for R2, positive differences favor SVR."
    )
    table = doc.add_table(rows=1, cols=len(manuscript.columns))
    table.style = "Table Grid"
    for i, col in enumerate(mancript_cols := list(manuscript.columns)):
        table.rows[0].cells[i].text = col
    for _, row in manuscript.iterrows():
        cells = table.add_row().cells
        for i, col in enumerate(mancript_cols):
            cells[i].text = str(row[col])
    doc.save(path)
    return True


def save_table_bundle(table: pd.DataFrame, manuscript: pd.DataFrame) -> dict[str, str]:
    paths: dict[str, str] = {}
    raw_csv = OUT / "TableS7_Paired_SVR_XGBoost_raw.csv"
    manuscript_csv = OUT / "TableS7_Paired_SVR_XGBoost_manuscript.csv"
    manuscript_md = OUT / "TableS7_Paired_SVR_XGBoost_manuscript.md"
    manuscript_tex = OUT / "TableS7_Paired_SVR_XGBoost_manuscript.tex"
    xlsx = OUT / "TableS7_Paired_SVR_XGBoost.xlsx"
    docx = OUT / "TableS7_Paired_SVR_XGBoost.docx"

    table.to_csv(raw_csv, index=False)
    manuscript.to_csv(manuscript_csv, index=False)
    manuscript_md.write_text(dataframe_to_markdown(manuscript), encoding="utf-8")
    manuscript_tex.write_text(dataframe_to_latex(manuscript), encoding="utf-8")
    paths["raw_csv"] = str(raw_csv)
    paths["manuscript_csv"] = str(manuscript_csv)
    paths["manuscript_md"] = str(manuscript_md)
    paths["manuscript_tex"] = str(manuscript_tex)

    try:
        write_xlsx(
            xlsx,
            {
                "Table S7 manuscript": manuscript,
                "raw statistics": table,
            },
        )
        paths["xlsx"] = str(xlsx)
    except Exception as exc:
        paths["xlsx_error"] = repr(exc)

    if write_docx_table(docx, manuscript):
        paths["docx"] = str(docx)
    else:
        paths["docx_error"] = "python-docx not available"

    return paths


def plot_forest(table: pd.DataFrame) -> list[str]:
    paths = []
    order = ["Overall", "cluster_c0.20", "cluster_c0.30", "cluster_c0.40", "scaffold", "stratified"]
    fig, axes = plt.subplots(2, 2, figsize=(11, 8.2), constrained_layout=True)
    for ax, metric in zip(axes.flat, METRIC_INFO):
        m = table[table["metric"] == metric].copy()
        m["comparison_set"] = pd.Categorical(m["comparison_set"], categories=order, ordered=True)
        m = m.sort_values("comparison_set")
        y = np.arange(len(m))
        x = m["paired_mean_difference_SVR_minus_XGBoost"].to_numpy(float)
        lo = m["bootstrap_95CI_low"].to_numpy(float)
        hi = m["bootstrap_95CI_high"].to_numpy(float)
        ax.errorbar(x, y, xerr=[x - lo, hi - x], fmt="o", color="black", ecolor="0.35", capsize=3)
        ax.axvline(0, color="0.45", lw=1, ls="--")
        ax.set_yticks(y, m["comparison_set"])
        ax.invert_yaxis()
        ax.set_title(METRIC_INFO[metric]["label"])
        ax.set_xlabel("Mean paired difference (SVR - XGBoost)")
        ax.grid(axis="x", color="0.9", lw=0.8)
    fig.suptitle(
        "Paired SVR vs XGBoost differences with bootstrap 95% confidence intervals",
        fontsize=13,
    )
    for ext in ("png", "svg", "pdf"):
        path = OUT / f"FigureS7_paired_mean_difference_forest.{ext}"
        fig.savefig(path, dpi=600 if ext == "png" else None, bbox_inches="tight")
        paths.append(str(path))
    plt.close(fig)
    return paths


def plot_difference_distributions(pair_level: pd.DataFrame) -> list[str]:
    paths = []
    overall = pair_level[pair_level["comparison_set"] == "Overall"].copy()
    metrics = list(METRIC_INFO.keys())
    fig, axes = plt.subplots(2, 2, figsize=(10.5, 7.6), constrained_layout=True)
    for ax, metric in zip(axes.flat, metrics):
        values = overall.loc[
            overall["metric"] == metric, "paired_difference_SVR_minus_XGBoost"
        ].to_numpy(float)
        ax.boxplot(values, vert=False, widths=0.45, showfliers=False)
        jitter = np.linspace(-0.08, 0.08, values.size)
        rng = np.random.default_rng(RNG_SEED + len(metric))
        rng.shuffle(jitter)
        ax.scatter(values, 1 + jitter, s=14, color="0.25", alpha=0.65)
        ax.axvline(0, color="0.45", lw=1, ls="--")
        ax.set_yticks([])
        ax.set_title(METRIC_INFO[metric]["label"])
        ax.set_xlabel("Paired difference (SVR - XGBoost)")
        ax.grid(axis="x", color="0.9", lw=0.8)
    fig.suptitle("Distribution of matched seed/fold paired differences", fontsize=13)
    for ext in ("png", "svg", "pdf"):
        path = OUT / f"FigureS8_paired_difference_distributions.{ext}"
        fig.savefig(path, dpi=600 if ext == "png" else None, bbox_inches="tight")
        paths.append(str(path))
    plt.close(fig)
    return paths


def main() -> None:
    OUT.mkdir(parents=True, exist_ok=True)
    results = pd.read_csv(RESULTS_CSV)
    wide = matched_pairs(results, MAIN_FRAC, MAIN_ALPHA)
    table, pair_level = paired_table(wide)
    manuscript = manuscript_view(table)
    alpha_sensitivity = alpha_width_sensitivity(results)

    table_paths = save_table_bundle(table, manuscript)

    pair_csv = OUT / "TableS7_pair_level_differences_alpha010_frac1.csv"
    pair_level.to_csv(pair_csv, index=False)

    alpha_csv = OUT / "TableS7_interval_width_alpha_sensitivity.csv"
    alpha_md = OUT / "TableS7_interval_width_alpha_sensitivity.md"
    alpha_sensitivity.to_csv(alpha_csv, index=False)
    alpha_view = alpha_sensitivity.copy()
    for col in [
        "SVR_width_mean",
        "XGBoost_width_mean",
        "paired_mean_difference_SVR_minus_XGBoost",
        "bootstrap_95CI_low",
        "bootstrap_95CI_high",
    ]:
        alpha_view[col] = alpha_view[col].map(lambda x: f"{x:.2f}")
    alpha_view["wilcoxon_p"] = alpha_sensitivity["wilcoxon_p"].map(p_format)
    alpha_md.write_text(dataframe_to_markdown(alpha_view), encoding="utf-8")

    forest_paths = plot_forest(table)
    distribution_paths = plot_difference_distributions(pair_level)

    with open(DATASET_META_JSON) as f:
        dataset_meta = json.load(f)
    with open(RUN_META_JSON) as f:
        run_meta = json.load(f)

    manifest = {
        "analysis": "Table S7 paired SVR vs XGBoost statistical comparison",
        "input_files": {
            "primary_results": str(RESULTS_CSV),
            "benchmark_summary_checked": str(SUMMARY_CSV),
            "dataset_metadata": str(DATASET_META_JSON),
            "run_metadata": str(RUN_META_JSON),
        },
        "dataset": dataset_meta,
        "run": run_meta,
        "settings": {
            "learning_fraction": MAIN_FRAC,
            "alpha": MAIN_ALPHA,
            "nominal_coverage": 1.0 - MAIN_ALPHA,
            "bootstrap_resamples": BOOTSTRAP_N,
            "bootstrap_seed": RNG_SEED,
            "paired_difference": "SVR minus XGBoost",
            "wilcoxon": "two-sided Wilcoxon signed-rank test on matched paired differences",
        },
        "matching_keys": ["regime", "seed", "fold", "frac", "alpha"],
        "n_matched_pairs_overall": int(wide.shape[0]),
        "n_matched_pairs_by_regime": wide.groupby("regime").size().to_dict(),
        "outputs": {
            **table_paths,
            "pair_level_csv": str(pair_csv),
            "alpha_sensitivity_csv": str(alpha_csv),
            "alpha_sensitivity_md": str(alpha_md),
            "figures": forest_paths + distribution_paths,
        },
    }
    manifest_path = OUT / "TableS7_analysis_manifest.json"
    manifest_path.write_text(json.dumps(manifest, indent=2), encoding="utf-8")

    notes = OUT / "TableS7_methods_and_file_notes.md"
    notes.write_text(
        "\n".join(
            [
                "# Table S7 Analysis Notes",
                "",
                "## Input files used",
                f"- `{RESULTS_CSV}`: primary seed/fold-level benchmark results used for all paired tests.",
                f"- `{SUMMARY_CSV}`: checked because the benchmark pipeline defines the manuscript summary at `frac = 1.0`.",
                f"- `{DATASET_META_JSON}`: records dataset identity, processed row count, target column, and SHA-256 hashes.",
                f"- `{RUN_META_JSON}`: records enabled models and benchmark output provenance.",
                "",
                "## Statistical choices",
                f"- Main comparison uses `frac = {MAIN_FRAC}` to match the benchmark's full-training summary.",
                f"- Conformal interval width uses `alpha = {MAIN_ALPHA:.2f}` (`{1-MAIN_ALPHA:.0%}` nominal coverage) for Table S7.",
                "- Matched units are identical `regime`, `seed`, `fold`, `frac`, and `alpha` rows for SVR and XGBoost.",
                "- Paired differences are `SVR - XGBoost`.",
                "- Wilcoxon signed-rank tests are two-sided.",
                f"- Bootstrap confidence intervals are percentile 95% CIs from {BOOTSTRAP_N:,} resamples of the paired differences.",
                "- For RMSE, MAE, and interval width, negative differences favor SVR. For R2, positive differences favor SVR.",
            ]
        )
        + "\n",
        encoding="utf-8",
    )

    print(json.dumps(manifest, indent=2))


if __name__ == "__main__":
    main()
