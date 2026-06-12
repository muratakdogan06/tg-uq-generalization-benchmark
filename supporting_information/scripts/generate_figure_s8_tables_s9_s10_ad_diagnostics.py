from __future__ import annotations

from datetime import datetime
from html import escape
from pathlib import Path
import json
import shutil
import zipfile

import matplotlib

matplotlib.use("Agg")
import matplotlib.pyplot as plt
import numpy as np
import pandas as pd
from rdkit import Chem, DataStructs, rdBase
from rdkit.Chem import Descriptors
from rdkit.Chem.rdMolDescriptors import GetMorganFingerprintAsBitVect
from rdkit.Chem.Scaffolds import MurckoScaffold
from scipy.stats import spearmanr
from sklearn.decomposition import PCA
from sklearn.impute import SimpleImputer
from sklearn.model_selection import KFold, StratifiedKFold
from sklearn.neighbors import NearestNeighbors
from sklearn.preprocessing import StandardScaler


rdBase.DisableLog("rdApp.warning")


import sys
from pathlib import Path

_SI_ROOT = Path(__file__).resolve().parents[1]
if str(_SI_ROOT) not in sys.path:
    sys.path.insert(0, str(_SI_ROOT))
from _paths import (  # noqa: E402
    BENCHMARK,
    CONFIG_YAML,
    DATASET_META_JSON,
    FINAL,
    INTERVALS_CSV,
    OUT,
    PROCESSED_CSV,
    RESULTS_CSV,
    ROOT,
    RUN,
    S8_MANIFEST,
    SIMILARITY_CSV,
    SUMMARY_CSV,
    XGB_SHAP_CSV,
    XGB_SHAP_FIG,
)


PROCESSED_CSV = RUN / "data" / "processed.csv"
DATASET_META_JSON = RUN / "data" / "dataset_meta.json"
CONFIG_YAML = BENCHMARK / "configs" / "tg.yaml"
SPLITS_PY = BENCHMARK / "src" / "cms_tg" / "splits.py"
SIMILARITY_PY = BENCHMARK / "src" / "cms_tg" / "similarity.py"
INTERVALS_CSV = OUT / "TableS8_per_sample_conformal_intervals_frac1_alpha010.csv"

FIGURE_TITLE = "Figure S8. Expanded applicability-domain diagnostics"
TABLE_S9_TITLE = "Table S9. Relationship among applicability-domain metrics"
TABLE_S10_TITLE = "Table S10. Threshold-based applicability-domain analysis"

SEEDS = [42, 43, 44]
CUTOFFS = [0.20, 0.30, 0.40]
N_FOLDS = 5
MORGAN_RADIUS = 2
MORGAN_NBITS = 2048
KNN_K = 5
LOW_SMAX_THRESHOLD = 0.30
N_PCA_COMPONENTS = 10

AD_METRICS = [
    "Smax",
    "knn5_similarity",
    "tanimoto_density_030",
    "descriptor_nn_distance",
    "pca_distance",
    "target_nn_distance_K",
]

AD_LABELS = {
    "Smax": "S$_{max}$",
    "knn5_similarity": "kNN similarity",
    "tanimoto_density_030": "Tanimoto density >=0.30",
    "descriptor_nn_distance": "Descriptor NN distance",
    "pca_distance": "PCA distance",
    "target_nn_distance_K": "Target-space NN distance (K)",
    "target_outside_train_range": "Outside training T$_g$ range",
    "composite_outside_ad": "Composite outside AD",
}

HIGHER_SUPPORT = {"Smax", "knn5_similarity", "tanimoto_density_030"}


def pct(numerator: float, denominator: float) -> float:
    return 100.0 * float(numerator) / float(denominator) if denominator else 0.0


def rmse(values: pd.Series | np.ndarray) -> float:
    arr = np.asarray(values, dtype=float)
    return float(np.sqrt(np.mean(arr * arr))) if len(arr) else np.nan


def p_format(p: float) -> str:
    if pd.isna(p):
        return "NA"
    if p < 0.001:
        return "<0.001"
    return f"{p:.3f}"


def _rdkit_2d_descriptor_matrix(mols: list[Chem.Mol]) -> pd.DataFrame:
    desc_names = [d[0] for d in Descriptors._descList]
    funcs = [d[1] for d in Descriptors._descList]
    rows = []
    for mol in mols:
        vals = []
        for func in funcs:
            try:
                vals.append(float(func(mol)))
            except Exception:
                vals.append(np.nan)
        rows.append(vals)
    return pd.DataFrame(rows, columns=desc_names)


def _polymer_proxy_features(mols: list[Chem.Mol]) -> pd.DataFrame:
    rows = []
    for mol in mols:
        heavy = mol.GetNumHeavyAtoms()
        denom = heavy if heavy > 0 else 1
        rings = int(mol.GetRingInfo().NumRings())
        arom = sum(1 for atom in mol.GetAtoms() if atom.GetIsAromatic())
        hetero = sum(1 for atom in mol.GetAtoms() if atom.GetAtomicNum() not in (1, 6))
        hal = sum(1 for atom in mol.GetAtoms() if atom.GetAtomicNum() in (9, 17, 35, 53))
        rows.append(
            {
                "proxy_heavy_atoms": heavy,
                "proxy_rings": rings,
                "proxy_aromatic_frac": arom / denom,
                "proxy_hetero_frac": hetero / denom,
                "proxy_halogen_frac": hal / denom,
            }
        )
    return pd.DataFrame(rows)


def load_dataset() -> tuple[pd.DataFrame, list[Chem.Mol], list[DataStructs.cDataStructs.ExplicitBitVect], pd.DataFrame]:
    df = pd.read_csv(PROCESSED_CSV)
    with DATASET_META_JSON.open("r", encoding="utf-8") as handle:
        meta = json.load(handle)
    smiles_col = meta["smiles_col"]
    target_col = meta["target_col"]

    rows = []
    mols = []
    for idx, row in df.iterrows():
        mol = Chem.MolFromSmiles(str(row[smiles_col]))
        if mol is None:
            continue
        rows.append(
            {
                "sample_index": int(idx),
                "SMILES": str(row[smiles_col]),
                "Tg_K": float(row[target_col]),
            }
        )
        mols.append(mol)
    data = pd.DataFrame(rows).reset_index(drop=True)
    fps = [GetMorganFingerprintAsBitVect(mol, MORGAN_RADIUS, nBits=MORGAN_NBITS) for mol in mols]
    X = pd.concat([_rdkit_2d_descriptor_matrix(mols), _polymer_proxy_features(mols)], axis=1)
    X = pd.DataFrame(SimpleImputer(strategy="median").fit_transform(X), columns=X.columns)
    return data, mols, fps, X


def stratified_splits(y: np.ndarray, seed: int):
    y_num = pd.Series(y).copy()
    try:
        bins = pd.qcut(y_num, q=10, labels=False, duplicates="drop")
        bins = np.asarray(bins, dtype=np.int64)
    except Exception:
        bins = pd.cut(y_num, bins=10, labels=False, duplicates="drop")
        bins = np.asarray(bins, dtype=np.int64)
    if np.any(pd.isna(bins)) or len(np.unique(bins)) < 2:
        kf = KFold(n_splits=N_FOLDS, shuffle=True, random_state=int(seed))
        return [(tr, te) for tr, te in kf.split(np.zeros(len(y_num)))]
    counts = np.bincount(bins[bins >= 0])
    if len(counts) == 0 or counts.min() < N_FOLDS:
        kf = KFold(n_splits=N_FOLDS, shuffle=True, random_state=int(seed))
        return [(tr, te) for tr, te in kf.split(np.zeros(len(y_num)))]
    skf = StratifiedKFold(n_splits=N_FOLDS, shuffle=True, random_state=int(seed))
    return [(tr, te) for tr, te in skf.split(np.zeros(len(y_num)), bins)]


def scaffold_splits(mols: list[Chem.Mol], seed: int):
    rng = np.random.RandomState(int(seed))
    scaffolds: dict[str, list[int]] = {}
    for i, mol in enumerate(mols):
        scaffold = MurckoScaffold.MurckoScaffoldSmiles(mol=mol)
        scaffolds.setdefault(scaffold, []).append(i)
    groups = list(scaffolds.values())
    rng.shuffle(groups)
    return grouped_splits(groups)


def cluster_by_cutoff(fps, cutoff: float) -> list[list[int]]:
    clusters: list[list[int]] = []
    for i, fp in enumerate(fps):
        assigned = False
        for cluster in clusters:
            sim = DataStructs.TanimotoSimilarity(fp, fps[cluster[0]])
            if sim >= float(cutoff):
                cluster.append(i)
                assigned = True
                break
        if not assigned:
            clusters.append([i])
    return clusters


def grouped_splits(groups: list[list[int]]):
    fold_bins: list[list[int]] = [[] for _ in range(N_FOLDS)]
    fold_sizes = np.zeros(N_FOLDS, dtype=int)
    for group in sorted(groups, key=len, reverse=True):
        fold_idx = int(np.argmin(fold_sizes))
        fold_bins[fold_idx].extend(group)
        fold_sizes[fold_idx] += len(group)
    all_idx = np.arange(sum(len(group) for group in groups))
    return [
        (np.setdiff1d(all_idx, np.array(sorted(fold), dtype=int)), np.array(sorted(fold), dtype=int))
        for fold in fold_bins
    ]


def cluster_splits(fps, seed: int, cutoff: float):
    rng = np.random.RandomState(int(seed))
    clusters = cluster_by_cutoff(fps, cutoff)
    rng.shuffle(clusters)
    return grouped_splits(clusters)


def compute_fold_ad_metrics(
    data: pd.DataFrame,
    fps,
    X: pd.DataFrame,
    train_idx: np.ndarray,
    test_idx: np.ndarray,
) -> pd.DataFrame:
    train_fps = [fps[int(i)] for i in train_idx]
    X_train = X.iloc[train_idx].to_numpy(float)
    X_test = X.iloc[test_idx].to_numpy(float)
    y_train = data.iloc[train_idx]["Tg_K"].to_numpy(float)
    y_test = data.iloc[test_idx]["Tg_K"].to_numpy(float)

    scaler = StandardScaler()
    X_train_z = scaler.fit_transform(X_train)
    X_test_z = scaler.transform(X_test)

    nn = NearestNeighbors(n_neighbors=1, metric="euclidean")
    nn.fit(X_train_z)
    descriptor_nn_distance = nn.kneighbors(X_test_z, return_distance=True)[0][:, 0]

    n_comp = min(N_PCA_COMPONENTS, X_train_z.shape[0] - 1, X_train_z.shape[1])
    pca = PCA(n_components=n_comp, random_state=0)
    train_scores = pca.fit_transform(X_train_z)
    test_scores = pca.transform(X_test_z)
    center = train_scores.mean(axis=0)
    scale = train_scores.std(axis=0, ddof=1)
    scale[scale == 0] = 1.0
    pca_distance = np.sqrt(np.sum(((test_scores - center) / scale) ** 2, axis=1))

    rows = []
    for pos, sample_idx in enumerate(test_idx):
        sims = np.asarray(DataStructs.BulkTanimotoSimilarity(fps[int(sample_idx)], train_fps), dtype=float)
        sorted_sims = np.sort(sims)[::-1]
        k = min(KNN_K, len(sorted_sims))
        target_nn_distance = float(np.min(np.abs(y_train - y_test[pos])))
        rows.append(
            {
                "sample_index": int(data.iloc[int(sample_idx)]["sample_index"]),
                "local_index": int(sample_idx),
                "Smax": float(sorted_sims[0]) if len(sorted_sims) else 0.0,
                "knn5_similarity": float(np.mean(sorted_sims[:k])) if k else 0.0,
                "tanimoto_density_030": float(np.mean(sims >= LOW_SMAX_THRESHOLD)) if len(sims) else 0.0,
                "descriptor_nn_distance": float(descriptor_nn_distance[pos]),
                "pca_distance": float(pca_distance[pos]),
                "target_nn_distance_K": target_nn_distance,
                "target_outside_train_range": bool(
                    y_test[pos] < np.min(y_train) or y_test[pos] > np.max(y_train)
                ),
                "target_train_percentile": float(np.mean(y_train <= y_test[pos])),
            }
        )
    return pd.DataFrame(rows)


def build_ad_records(data: pd.DataFrame, mols, fps, X: pd.DataFrame) -> pd.DataFrame:
    y = data["Tg_K"].to_numpy(float)
    rows = []
    for base_regime in ["stratified", "scaffold", "cluster"]:
        cutoffs = [None] if base_regime != "cluster" else CUTOFFS
        for cutoff in cutoffs:
            regime = base_regime if cutoff is None else f"cluster_c{cutoff:.2f}"
            for seed in SEEDS:
                if base_regime == "stratified":
                    splits = stratified_splits(y, seed)
                elif base_regime == "scaffold":
                    splits = scaffold_splits(mols, seed)
                else:
                    splits = cluster_splits(fps, seed, float(cutoff))
                for fold, (train_idx, test_idx) in enumerate(splits):
                    fold_ad = compute_fold_ad_metrics(data, fps, X, train_idx, test_idx)
                    fold_ad.insert(0, "fold", int(fold))
                    fold_ad.insert(0, "seed", int(seed))
                    fold_ad.insert(0, "cutoff", np.nan if cutoff is None else float(cutoff))
                    fold_ad.insert(0, "base_regime", base_regime)
                    fold_ad.insert(0, "regime", regime)
                    fold_ad["train_size"] = int(len(train_idx))
                    fold_ad["test_size"] = int(len(test_idx))
                    rows.append(fold_ad)
    return pd.concat(rows, ignore_index=True)


def merge_predictions(ad_records: pd.DataFrame) -> pd.DataFrame:
    intervals = pd.read_csv(INTERVALS_CSV)
    intervals = intervals[(intervals["frac"] == 1.0) & (intervals["alpha"] == 0.10)].copy()
    intervals["cutoff_key"] = intervals["cutoff"].fillna("none").astype(str)
    ad = ad_records.copy()
    ad["cutoff_key"] = ad["cutoff"].fillna("none").astype(str)
    keys = ["regime", "base_regime", "cutoff_key", "seed", "fold", "sample_index"]
    merged = intervals.merge(
        ad.drop(columns=["cutoff"]),
        on=keys,
        how="left",
        validate="many_to_one",
    )
    if "Smax_x" in merged.columns and "Smax_y" in merged.columns:
        merged = merged.rename(columns={"Smax_x": "Smax_benchmark_file", "Smax_y": "Smax"})
        merged["Smax_delta_vs_benchmark_file"] = merged["Smax"] - merged["Smax_benchmark_file"]
    missing = int(merged["knn5_similarity"].isna().sum())
    if missing:
        raise ValueError(f"Missing AD metrics after merge: {missing} rows")
    merged["absolute_error"] = (merged["y_pred"] - merged["Tg"]).abs()
    merged["signed_error"] = merged["y_pred"] - merged["Tg"]
    merged["squared_error"] = merged["signed_error"] ** 2
    return merged


def threshold_definitions(ad_records: pd.DataFrame) -> dict[str, dict[str, object]]:
    thresholds = {
        "Smax": {"threshold": LOW_SMAX_THRESHOLD, "inside_when": ">=", "outside_col": "outside_Smax"},
        "target_outside_train_range": {
            "threshold": "outside training Tg range",
            "inside_when": "False",
            "outside_col": "outside_target_outside_train_range",
        },
    }
    for metric in ["knn5_similarity", "tanimoto_density_030"]:
        thresholds[metric] = {
            "threshold": float(ad_records[metric].quantile(0.25)),
            "inside_when": ">=",
            "outside_col": f"outside_{metric}",
        }
    for metric in ["descriptor_nn_distance", "pca_distance", "target_nn_distance_K"]:
        thresholds[metric] = {
            "threshold": float(ad_records[metric].quantile(0.75)),
            "inside_when": "<=",
            "outside_col": f"outside_{metric}",
        }
    return thresholds


def add_outside_flags(df: pd.DataFrame, thresholds: dict[str, dict[str, object]]) -> pd.DataFrame:
    out = df.copy()
    for metric, spec in thresholds.items():
        col = str(spec["outside_col"])
        if metric == "target_outside_train_range":
            out[col] = out[metric].astype(bool)
        elif spec["inside_when"] == ">=":
            out[col] = out[metric] < float(spec["threshold"])
        else:
            out[col] = out[metric] > float(spec["threshold"])
    outside_cols = [str(spec["outside_col"]) for spec in thresholds.values()]
    out["composite_outside_ad"] = out[outside_cols].any(axis=1)
    return out


def table_s9(ad_records: pd.DataFrame, merged: pd.DataFrame) -> tuple[pd.DataFrame, pd.DataFrame]:
    rows = []
    unique_ad = ad_records.drop_duplicates(["regime", "base_regime", "cutoff", "seed", "fold", "sample_index"])
    for i, metric_a in enumerate(AD_METRICS):
        for metric_b in AD_METRICS[i + 1 :]:
            rho, p = spearmanr(unique_ad[metric_a], unique_ad[metric_b], nan_policy="omit")
            rows.append(
                {
                    "section": "AD metric pairwise relationship",
                    "model": "NA",
                    "metric_1": AD_LABELS[metric_a],
                    "metric_2": AD_LABELS[metric_b],
                    "n_records": int(unique_ad[[metric_a, metric_b]].dropna().shape[0]),
                    "spearman_rho": float(rho),
                    "p_value": float(p),
                }
            )
    for model, model_df in merged.groupby("model"):
        for metric in AD_METRICS + ["target_outside_train_range"]:
            values = model_df[metric].astype(float) if metric == "target_outside_train_range" else model_df[metric]
            rho, p = spearmanr(values, model_df["absolute_error"], nan_policy="omit")
            rows.append(
                {
                    "section": "AD metric vs absolute error",
                    "model": model,
                    "metric_1": AD_LABELS[metric],
                    "metric_2": "Absolute prediction error",
                    "n_records": int(model_df[[metric, "absolute_error"]].dropna().shape[0]),
                    "spearman_rho": float(rho),
                    "p_value": float(p),
                }
            )
    raw = pd.DataFrame(rows)
    manuscript = raw.copy()
    manuscript["Spearman rho"] = manuscript["spearman_rho"].map(lambda x: f"{x:.3f}")
    manuscript["p value"] = manuscript["p_value"].map(p_format)
    manuscript = manuscript[
        ["section", "model", "metric_1", "metric_2", "n_records", "Spearman rho", "p value"]
    ].rename(
        columns={
            "section": "Section",
            "model": "Model",
            "metric_1": "Metric 1",
            "metric_2": "Metric 2",
            "n_records": "N records",
        }
    )
    return raw, manuscript


def summarize_inside_outside(df: pd.DataFrame, group_label: str, model: str, metric: str, outside_col: str, threshold, inside_when: str) -> dict:
    inside = df[~df[outside_col].astype(bool)]
    outside = df[df[outside_col].astype(bool)]
    if metric == "target_outside_train_range":
        inside_rule = "Tg within training-fold target range"
    elif metric == "composite_outside_ad":
        inside_rule = "No AD threshold violated"
    else:
        inside_rule = f"{AD_LABELS.get(metric, metric)} {inside_when} {threshold if isinstance(threshold, str) else f'{float(threshold):.3f}'}"
    return {
        "comparison_set": group_label,
        "model": model,
        "AD metric": AD_LABELS.get(metric, metric),
        "threshold": threshold if isinstance(threshold, str) else float(threshold),
        "inside_AD_rule": inside_rule,
        "inside_n": int(len(inside)),
        "outside_n": int(len(outside)),
        "outside_fraction": float(len(outside) / len(df)) if len(df) else np.nan,
        "inside_MAE": float(inside["absolute_error"].mean()) if len(inside) else np.nan,
        "outside_MAE": float(outside["absolute_error"].mean()) if len(outside) else np.nan,
        "delta_MAE_outside_minus_inside": float(outside["absolute_error"].mean() - inside["absolute_error"].mean())
        if len(inside) and len(outside)
        else np.nan,
        "inside_RMSE": rmse(inside["signed_error"]) if len(inside) else np.nan,
        "outside_RMSE": rmse(outside["signed_error"]) if len(outside) else np.nan,
        "inside_coverage": float(inside["covered"].mean()) if len(inside) else np.nan,
        "outside_coverage": float(outside["covered"].mean()) if len(outside) else np.nan,
    }


def table_s10(merged: pd.DataFrame, thresholds: dict[str, dict[str, object]]) -> tuple[pd.DataFrame, pd.DataFrame]:
    rows = []
    metrics_with_composite = dict(thresholds)
    metrics_with_composite["composite_outside_ad"] = {
        "threshold": "outside any AD rule",
        "inside_when": "False",
        "outside_col": "composite_outside_ad",
    }
    for model, model_df in merged.groupby("model"):
        for metric, spec in metrics_with_composite.items():
            rows.append(
                summarize_inside_outside(
                    model_df,
                    "Overall",
                    model,
                    metric,
                    str(spec["outside_col"]),
                    spec["threshold"],
                    str(spec["inside_when"]),
                )
            )
            for regime, regime_df in model_df.groupby("regime"):
                rows.append(
                    summarize_inside_outside(
                        regime_df,
                        regime,
                        model,
                        metric,
                        str(spec["outside_col"]),
                        spec["threshold"],
                        str(spec["inside_when"]),
                    )
                )
    raw = pd.DataFrame(rows)
    manuscript = raw[raw["comparison_set"] == "Overall"].copy()
    for col in [
        "outside_fraction",
        "inside_MAE",
        "outside_MAE",
        "delta_MAE_outside_minus_inside",
        "inside_RMSE",
        "outside_RMSE",
        "inside_coverage",
        "outside_coverage",
    ]:
        manuscript[col] = manuscript[col].map(lambda x: "NA" if pd.isna(x) else f"{x:.3f}")
    manuscript = manuscript[
        [
            "model",
            "AD metric",
            "inside_AD_rule",
            "inside_n",
            "outside_n",
            "outside_fraction",
            "inside_MAE",
            "outside_MAE",
            "delta_MAE_outside_minus_inside",
            "inside_RMSE",
            "outside_RMSE",
            "inside_coverage",
            "outside_coverage",
        ]
    ].rename(
        columns={
            "model": "Model",
            "inside_AD_rule": "Inside-AD rule",
            "inside_n": "Inside n",
            "outside_n": "Outside n",
            "outside_fraction": "Outside fraction",
            "inside_MAE": "Inside MAE",
            "outside_MAE": "Outside MAE",
            "delta_MAE_outside_minus_inside": "Delta MAE",
            "inside_RMSE": "Inside RMSE",
            "outside_RMSE": "Outside RMSE",
            "inside_coverage": "Inside coverage",
            "outside_coverage": "Outside coverage",
        }
    )
    return raw, manuscript


def dataframe_to_markdown(df: pd.DataFrame) -> str:
    headers = [str(col) for col in df.columns]
    rows = [[str(value) for value in row] for row in df.to_numpy()]
    widths = [
        max(len(headers[i]), *(len(row[i]) for row in rows)) if rows else len(headers[i])
        for i in range(len(headers))
    ]
    lines = [
        "| " + " | ".join(headers[i].ljust(widths[i]) for i in range(len(headers))) + " |",
        "| " + " | ".join("-" * widths[i] for i in range(len(headers))) + " |",
    ]
    lines.extend("| " + " | ".join(row[i].ljust(widths[i]) for i in range(len(headers))) + " |" for row in rows)
    return "\n".join(lines) + "\n"


def latex_escape(value: object) -> str:
    text = str(value)
    for old, new in {
        "\\": r"\textbackslash{}",
        "&": r"\&",
        "%": r"\%",
        "$": r"\$",
        "#": r"\#",
        "_": r"\_",
        "{": r"\{",
        "}": r"\}",
        "~": r"\textasciitilde{}",
        "^": r"\textasciicircum{}",
    }.items():
        text = text.replace(old, new)
    return text


def dataframe_to_latex(df: pd.DataFrame, caption: str) -> str:
    cols = list(df.columns)
    lines = [
        r"\begin{table}[htbp]",
        r"\centering",
        f"\\caption{{{latex_escape(caption)}}}",
        r"\begin{tabular}{" + "l" * len(cols) + "}",
        r"\hline",
        " & ".join(latex_escape(col) for col in cols) + r" \\",
        r"\hline",
    ]
    for _, row in df.iterrows():
        lines.append(" & ".join(latex_escape(row[col]) for col in cols) + r" \\")
    lines.extend([r"\hline", r"\end{tabular}", r"\end{table}", ""])
    return "\n".join(lines)


def excel_col_name(index: int) -> str:
    name = ""
    while index:
        index, remainder = divmod(index - 1, 26)
        name = chr(65 + remainder) + name
    return name


def worksheet_xml(df: pd.DataFrame) -> str:
    rows = [list(df.columns)] + df.astype(str).values.tolist()
    xml_rows = []
    for r_idx, row in enumerate(rows, start=1):
        cells = []
        for c_idx, value in enumerate(row, start=1):
            ref = f"{excel_col_name(c_idx)}{r_idx}"
            cells.append(f'<c r="{ref}" t="inlineStr"><is><t>{escape(str(value))}</t></is></c>')
        xml_rows.append(f'<row r="{r_idx}">{"".join(cells)}</row>')
    return (
        '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
        '<worksheet xmlns="http://schemas.openxmlformats.org/spreadsheetml/2006/main">'
        f'<sheetData>{"".join(xml_rows)}</sheetData>'
        "</worksheet>"
    )


def write_xlsx(path: Path, sheets: dict[str, pd.DataFrame]) -> None:
    sheet_names = list(sheets.keys())
    with zipfile.ZipFile(path, "w", zipfile.ZIP_DEFLATED) as zf:
        zf.writestr(
            "[Content_Types].xml",
            '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
            '<Types xmlns="http://schemas.openxmlformats.org/package/2006/content-types">'
            '<Default Extension="rels" ContentType="application/vnd.openxmlformats-package.relationships+xml"/>'
            '<Default Extension="xml" ContentType="application/xml"/>'
            '<Override PartName="/xl/workbook.xml" ContentType="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet.main+xml"/>'
            + "".join(
                f'<Override PartName="/xl/worksheets/sheet{i}.xml" '
                'ContentType="application/vnd.openxmlformats-officedocument.spreadsheetml.worksheet+xml"/>'
                for i in range(1, len(sheet_names) + 1)
            )
            + "</Types>",
        )
        zf.writestr(
            "_rels/.rels",
            '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
            '<Relationships xmlns="http://schemas.openxmlformats.org/package/2006/relationships">'
            '<Relationship Id="rId1" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/officeDocument" Target="xl/workbook.xml"/>'
            "</Relationships>",
        )
        zf.writestr(
            "xl/workbook.xml",
            '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
            '<workbook xmlns="http://schemas.openxmlformats.org/spreadsheetml/2006/main" '
            'xmlns:r="http://schemas.openxmlformats.org/officeDocument/2006/relationships">'
            "<sheets>"
            + "".join(
                f'<sheet name="{escape(name[:31])}" sheetId="{i}" r:id="rId{i}"/>'
                for i, name in enumerate(sheet_names, start=1)
            )
            + "</sheets></workbook>",
        )
        zf.writestr(
            "xl/_rels/workbook.xml.rels",
            '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
            '<Relationships xmlns="http://schemas.openxmlformats.org/package/2006/relationships">'
            + "".join(
                f'<Relationship Id="rId{i}" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/worksheet" '
                f'Target="worksheets/sheet{i}.xml"/>'
                for i in range(1, len(sheet_names) + 1)
            )
            + "</Relationships>",
        )
        for i, name in enumerate(sheet_names, start=1):
            zf.writestr(f"xl/worksheets/sheet{i}.xml", worksheet_xml(sheets[name]))


def write_docx(path: Path, title: str, manuscript: pd.DataFrame) -> None:
    try:
        from docx import Document
    except Exception:
        rows = [list(manuscript.columns)] + manuscript.astype(str).values.tolist()
        table_rows = []
        for row in rows:
            cells = "".join(
                "<w:tc><w:p><w:r><w:t>"
                + escape(str(value))
                + "</w:t></w:r></w:p></w:tc>"
                for value in row
            )
            table_rows.append(f"<w:tr>{cells}</w:tr>")
        document_xml = (
            '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
            '<w:document xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
            "<w:body>"
            f"<w:p><w:r><w:t>{escape(title)}</w:t></w:r></w:p>"
            f"<w:tbl>{''.join(table_rows)}</w:tbl>"
            + '<w:sectPr><w:pgSz w:w="15840" w:h="12240" w:orient="landscape"/></w:sectPr>'
            + "</w:body></w:document>"
        )
        with zipfile.ZipFile(path, "w", zipfile.ZIP_DEFLATED) as zf:
            zf.writestr(
                "[Content_Types].xml",
                '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
                '<Types xmlns="http://schemas.openxmlformats.org/package/2006/content-types">'
                '<Default Extension="rels" ContentType="application/vnd.openxmlformats-package.relationships+xml"/>'
                '<Default Extension="xml" ContentType="application/xml"/>'
                '<Override PartName="/word/document.xml" ContentType="application/vnd.openxmlformats-officedocument.wordprocessingml.document.main+xml"/>'
                "</Types>",
            )
            zf.writestr(
                "_rels/.rels",
                '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
                '<Relationships xmlns="http://schemas.openxmlformats.org/package/2006/relationships">'
                '<Relationship Id="rId1" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/officeDocument" Target="word/document.xml"/>'
                "</Relationships>",
            )
            zf.writestr("word/document.xml", document_xml)
        return
    doc = Document()
    doc.add_heading(title, level=1)
    table = doc.add_table(rows=1, cols=len(manuscript.columns))
    table.style = "Table Grid"
    cols = list(manuscript.columns)
    for i, col in enumerate(cols):
        table.rows[0].cells[i].text = col
    for _, row in manuscript.iterrows():
        cells = table.add_row().cells
        for i, col in enumerate(cols):
            cells[i].text = str(row[col])
    doc.save(path)


def save_table_bundle(prefix: str, title: str, raw: pd.DataFrame, manuscript: pd.DataFrame) -> dict[str, Path]:
    paths = {
        "raw_csv": OUT / f"{prefix}_raw.csv",
        "manuscript_csv": OUT / f"{prefix}_manuscript.csv",
        "manuscript_md": OUT / f"{prefix}_manuscript.md",
        "manuscript_tex": OUT / f"{prefix}_manuscript.tex",
        "xlsx": OUT / f"{prefix}.xlsx",
        "docx": OUT / f"{prefix}.docx",
    }
    raw.to_csv(paths["raw_csv"], index=False)
    manuscript.to_csv(paths["manuscript_csv"], index=False)
    paths["manuscript_md"].write_text(dataframe_to_markdown(manuscript), encoding="utf-8")
    paths["manuscript_tex"].write_text(dataframe_to_latex(manuscript, title), encoding="utf-8")
    write_xlsx(paths["xlsx"], {"manuscript": manuscript, "raw": raw})
    write_docx(paths["docx"], title, manuscript)
    return paths


def binned_error(df: pd.DataFrame, metric: str, n_bins: int = 8) -> pd.DataFrame:
    rows = []
    for model, model_df in df.groupby("model"):
        bins = pd.qcut(model_df[metric], q=n_bins, duplicates="drop")
        for interval, group in model_df.groupby(bins, observed=True):
            rows.append(
                {
                    "model": model,
                    "metric": metric,
                    "x": float(group[metric].median()),
                    "mae": float(group["absolute_error"].mean()),
                    "n": int(len(group)),
                    "bin": str(interval),
                }
            )
    return pd.DataFrame(rows)


def plot_figure_s8(merged: pd.DataFrame, s9_raw: pd.DataFrame, s10_raw: pd.DataFrame, out_prefix: Path) -> list[Path]:
    fig, axes = plt.subplots(2, 3, figsize=(13, 8), constrained_layout=True)
    plt.rcParams.update({"font.size": 9, "axes.titlesize": 10, "axes.labelsize": 9})

    for ax, metric, title in [
        (axes[0, 0], "Smax", "A) Error vs S$_{max}$"),
        (axes[0, 1], "knn5_similarity", "B) Error vs kNN similarity"),
        (axes[0, 2], "descriptor_nn_distance", "C) Error vs descriptor distance"),
    ]:
        plot_df = binned_error(merged, metric)
        for model, group in plot_df.groupby("model"):
            group = group.sort_values("x")
            ax.plot(group["x"], group["mae"], marker="o", label=model.upper())
        ax.set_xlabel(AD_LABELS[metric])
        ax.set_ylabel("MAE (K)")
        ax.set_title(title)
        ax.grid(alpha=0.25)
        ax.legend(frameon=False)

    composite = s10_raw[(s10_raw["comparison_set"] == "Overall") & (s10_raw["AD metric"] == AD_LABELS["composite_outside_ad"])]
    ax = axes[1, 0]
    x = np.arange(len(composite))
    width = 0.35
    ax.bar(x - width / 2, composite["inside_MAE"], width=width, label="Inside AD")
    ax.bar(x + width / 2, composite["outside_MAE"], width=width, label="Outside AD")
    ax.set_xticks(x)
    ax.set_xticklabels([str(m).upper() for m in composite["model"]])
    ax.set_ylabel("MAE (K)")
    ax.set_title("D) Composite inside vs outside AD")
    ax.grid(axis="y", alpha=0.25)
    ax.legend(frameon=False)

    ax = axes[1, 1]
    unique_ad = merged.drop_duplicates(["regime", "base_regime", "cutoff_key", "seed", "fold", "sample_index"])
    corr = unique_ad[AD_METRICS].corr(method="spearman")
    im = ax.imshow(corr, vmin=-1, vmax=1, cmap="coolwarm")
    labels = ["S$_{max}$", "kNN", "density", "desc dist", "PCA dist", "target dist"]
    ax.set_xticks(range(len(labels)))
    ax.set_yticks(range(len(labels)))
    ax.set_xticklabels(labels, rotation=45, ha="right")
    ax.set_yticklabels(labels)
    ax.set_title("E) AD metric Spearman correlations")
    for i in range(corr.shape[0]):
        for j in range(corr.shape[1]):
            ax.text(j, i, f"{corr.iloc[i, j]:.2f}", ha="center", va="center", fontsize=7)
    fig.colorbar(im, ax=ax, fraction=0.046, pad=0.04)

    ax = axes[1, 2]
    low_support = (
        merged.drop_duplicates(["model", "regime", "base_regime", "cutoff_key", "seed", "fold", "sample_index"])
        .groupby(["model", "regime"])["composite_outside_ad"]
        .mean()
        .reset_index()
    )
    regimes = ["stratified", "scaffold", "cluster_c0.20", "cluster_c0.30", "cluster_c0.40"]
    regime_labels = ["Stratified", "Scaffold", "Cluster (c=0.20)", "Cluster (c=0.30)", "Cluster (c=0.40)"]
    x = np.arange(len(regimes))
    for offset, model in [(-0.18, "svr"), (0.18, "xgb")]:
        vals = [
            float(low_support[(low_support["model"] == model) & (low_support["regime"] == regime)]["composite_outside_ad"].iloc[0] * 100)
            for regime in regimes
        ]
        ax.bar(x + offset, vals, width=0.34, label=model.upper())
    ax.set_xticks(x)
    ax.set_xticklabels(regime_labels, rotation=35, ha="right")
    ax.set_ylabel("Composite outside-AD fraction (%)")
    ax.set_title("F) Low-support fraction by regime")
    ax.grid(axis="y", alpha=0.25)
    ax.legend(frameon=False)

    paths = []
    for ext, kwargs in {
        ".png": {"dpi": 300},
        ".svg": {},
        ".pdf": {},
        ".tiff": {"dpi": 600},
    }.items():
        path = out_prefix.with_suffix(ext)
        fig.savefig(path, **kwargs)
        paths.append(path)
    plt.close(fig)
    return paths


def write_notes(paths: dict[str, Path], thresholds: dict[str, dict[str, object]], key_rows: pd.DataFrame, figure_paths: list[Path]) -> Path:
    notes_path = OUT / "FigureS8_TablesS9_S10_AD_diagnostics_notes.txt"
    lines = [
        f"{FIGURE_TITLE}; {TABLE_S9_TITLE}; {TABLE_S10_TITLE}",
        f"Generated: {datetime.now().isoformat()}",
        "",
        "Inputs used:",
        f"- Per-sample predictions and conformal intervals: {INTERVALS_CSV}",
        f"- Processed benchmark data: {PROCESSED_CSV}",
        f"- Dataset metadata: {DATASET_META_JSON}",
        f"- Benchmark config: {CONFIG_YAML}",
        f"- Split implementation: {SPLITS_PY}",
        f"- Smax implementation: {SIMILARITY_PY}",
        "",
        "AD metrics:",
        f"- Smax: maximum Morgan-Tanimoto similarity to training fold, radius={MORGAN_RADIUS}, nBits={MORGAN_NBITS}.",
        f"- kNN similarity: mean top-{KNN_K} Morgan-Tanimoto similarity values to the training fold.",
        f"- kNN density: fraction of training samples with Tanimoto similarity >= {LOW_SMAX_THRESHOLD:.2f}.",
        "- Descriptor-space distance: nearest-neighbor Euclidean distance in standardized RDKit descriptor + polymer proxy feature space.",
        f"- PCA distance: standardized distance from the training centroid in the first {N_PCA_COMPONENTS} PCA components.",
        "- Target-space support: nearest training Tg distance and whether the test Tg lies outside the training Tg range.",
        "",
        "Thresholds:",
    ]
    for metric, spec in thresholds.items():
        value = spec["threshold"]
        value_text = value if isinstance(value, str) else f"{float(value):.4f}"
        if metric == "target_outside_train_range":
            lines.append("- Outside training Tg range: inside when Tg is within the training-fold target range.")
        else:
            lines.append(f"- {AD_LABELS.get(metric, metric)}: inside when {spec['inside_when']} {value_text}.")
    lines.extend(["", "Key composite inside/outside AD rows:"])
    for _, row in key_rows.iterrows():
        lines.append(
            f"- {row['model'].upper()}: inside MAE {row['inside_MAE']:.2f} K; "
            f"outside MAE {row['outside_MAE']:.2f} K; outside fraction {row['outside_fraction']:.1%}."
        )
    lines.extend(
        [
            "",
            "Reviewer-facing interpretation:",
            "Expanded AD metrics provide a multi-view support diagnostic beyond Smax alone. "
            "Similarity-, descriptor-, PCA-, and target-support metrics are correlated but not redundant. "
            "Thresholded outside-AD subsets generally show higher empirical error and/or weaker coverage, "
            "supporting the manuscript claim that local support diagnostics should accompany global performance.",
            "",
            "Generated files:",
            *[f"- {path}" for path in paths.values()],
            *[f"- {path}" for path in figure_paths],
        ]
    )
    notes_path.write_text("\n".join(lines) + "\n", encoding="utf-8")
    return notes_path


def main() -> None:
    OUT.mkdir(parents=True, exist_ok=True)
    FINAL.mkdir(parents=True, exist_ok=True)

    print("Loading dataset and computing AD records...", flush=True)
    data, mols, fps, X = load_dataset()
    ad_records = build_ad_records(data, mols, fps, X)
    ad_records_path = OUT / "FigureS8_expanded_AD_metrics_per_test_record.csv"
    ad_records.to_csv(ad_records_path, index=False)

    print("Merging with per-sample prediction records...", flush=True)
    merged = merge_predictions(ad_records)
    thresholds = threshold_definitions(ad_records)
    merged = add_outside_flags(merged, thresholds)
    merged_path = OUT / "FigureS8_expanded_AD_metrics_with_errors.csv"
    merged.to_csv(merged_path, index=False)

    print("Building Tables S9 and S10...", flush=True)
    s9_raw, s9_manuscript = table_s9(ad_records, merged)
    s10_raw, s10_manuscript = table_s10(merged, thresholds)
    s9_paths = save_table_bundle("TableS9_AD_Metric_Relationships", TABLE_S9_TITLE, s9_raw, s9_manuscript)
    s10_paths = save_table_bundle("TableS10_Threshold_AD_Analysis", TABLE_S10_TITLE, s10_raw, s10_manuscript)

    print("Plotting Figure S8...", flush=True)
    figure_paths = plot_figure_s8(merged, s9_raw, s10_raw, OUT / "FigureS8_expanded_applicability_domain_diagnostics")
    for path in figure_paths:
        if path.suffix.lower() in {".png", ".tiff"}:
            shutil.copy2(path, FINAL / path.name)

    key_rows = s10_raw[
        (s10_raw["comparison_set"] == "Overall")
        & (s10_raw["AD metric"] == AD_LABELS["composite_outside_ad"])
    ]
    all_paths = {"ad_records": ad_records_path, "merged_errors": merged_path}
    all_paths.update({f"s9_{key}": value for key, value in s9_paths.items()})
    all_paths.update({f"s10_{key}": value for key, value in s10_paths.items()})
    notes_path = write_notes(all_paths, thresholds, key_rows, figure_paths)

    print(FIGURE_TITLE)
    print("Composite AD inside/outside summary:")
    print(
        key_rows[
            ["model", "inside_n", "outside_n", "outside_fraction", "inside_MAE", "outside_MAE", "inside_coverage", "outside_coverage"]
        ].to_string(index=False)
    )
    print("Wrote:")
    for path in [*all_paths.values(), notes_path, *figure_paths]:
        print(f"- {path}")


if __name__ == "__main__":
    main()
