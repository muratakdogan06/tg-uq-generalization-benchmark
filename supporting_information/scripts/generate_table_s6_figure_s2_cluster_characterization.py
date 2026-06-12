from __future__ import annotations

from datetime import datetime
from html import escape
from pathlib import Path
import json
import shutil
import zipfile

import matplotlib

matplotlib.use("Agg")
import matplotlib.pyplot as plt
import numpy as np
import pandas as pd
from rdkit import Chem, DataStructs, rdBase
from rdkit.Chem.rdMolDescriptors import GetMorganFingerprintAsBitVect

rdBase.DisableLog("rdApp.warning")


import sys
from pathlib import Path

_SI_ROOT = Path(__file__).resolve().parents[1]
if str(_SI_ROOT) not in sys.path:
    sys.path.insert(0, str(_SI_ROOT))
from _paths import (  # noqa: E402
    BENCHMARK,
    CONFIG_YAML,
    DATASET_META_JSON,
    FINAL,
    INTERVALS_CSV,
    OUT,
    PROCESSED_CSV,
    RESULTS_CSV,
    ROOT,
    RUN,
    S8_MANIFEST,
    SIMILARITY_CSV,
    SUMMARY_CSV,
    XGB_SHAP_CSV,
    XGB_SHAP_FIG,
)


PROCESSED_CSV = RUN / "data" / "processed.csv"
DATASET_META_JSON = RUN / "data" / "dataset_meta.json"
CONFIG_YAML = BENCHMARK / "configs" / "tg.yaml"
SPLITS_PY = BENCHMARK / "src" / "cms_tg" / "splits.py"
SIMILARITY_PY = BENCHMARK / "src" / "cms_tg" / "similarity.py"

TABLE_TITLE = "Table S6. Cluster-split characterization across Morgan-Tanimoto cutoffs"
FIGURE_TITLE = "Figure S2. Fold-level Tg and chemistry-distribution balance under cluster-based splitting"

SEEDS = [42, 43, 44]
CUTOFFS = [0.20, 0.30, 0.40]
N_FOLDS = 5
MORGAN_RADIUS = 2
MORGAN_NBITS = 2048
LOW_SMAX_THRESHOLD = 0.30
SEED_FOLD_TEST_SETS = len(SEEDS) * N_FOLDS


def table_s6_caption(hetero_threshold: float) -> str:
    return (
        f"{TABLE_TITLE} Cluster splits use Morgan fingerprints (radius {MORGAN_RADIUS}, "
        f"{MORGAN_NBITS} bits) and leader-style Tanimoto grouping, matching the benchmark "
        "implementation. Smax is the maximum Morgan-Tanimoto similarity from each test repeat "
        "unit to its training fold. Chemistry fold fractions summarize "
        f"{SEED_FOLD_TEST_SETS} seed-fold test sets per cutoff; heteroatom-rich is defined at "
        f"the dataset median heteroatom fraction ({hetero_threshold:.3f}). Low-Smax tail counts "
        "n (%) use all test repeat units pooled across those seed-fold test sets as the "
        "denominator. Cutoffs are operational stress-test settings, not universal polymer OOD "
        "thresholds."
    )


def pct(numerator: float, denominator: float) -> float:
    return 100.0 * float(numerator) / float(denominator) if denominator else 0.0


def q25(values: pd.Series | np.ndarray) -> float:
    return float(np.percentile(np.asarray(values, dtype=float), 25))


def q75(values: pd.Series | np.ndarray) -> float:
    return float(np.percentile(np.asarray(values, dtype=float), 75))


def fmt_iqr(values: pd.Series | np.ndarray, decimals: int = 2) -> str:
    arr = np.asarray(values, dtype=float)
    return f"{np.median(arr):.{decimals}f} [{np.percentile(arr, 25):.{decimals}f}, {np.percentile(arr, 75):.{decimals}f}]"


def fmt_range(values: pd.Series | np.ndarray, decimals: int = 1) -> str:
    arr = np.asarray(values, dtype=float)
    return f"{np.min(arr):.{decimals}f}-{np.max(arr):.{decimals}f}"


def size_distribution_text(sizes: pd.Series | np.ndarray) -> str:
    arr = np.asarray(sizes, dtype=int)
    bins = [
        ("1", arr == 1),
        ("2", arr == 2),
        ("3-5", (arr >= 3) & (arr <= 5)),
        ("6-10", (arr >= 6) & (arr <= 10)),
        (">10", arr > 10),
    ]
    return "; ".join(f"{label}: {int(mask.sum())}" for label, mask in bins)


def load_dataset() -> tuple[pd.DataFrame, list[Chem.Mol], list[DataStructs.cDataStructs.ExplicitBitVect], float]:
    df = pd.read_csv(PROCESSED_CSV)
    with DATASET_META_JSON.open("r", encoding="utf-8") as handle:
        meta = json.load(handle)
    smiles_col = meta["smiles_col"]
    target_col = meta["target_col"]

    rows = []
    mols = []
    for idx, row in df.iterrows():
        mol = Chem.MolFromSmiles(str(row[smiles_col]))
        if mol is None:
            continue
        heavy = mol.GetNumHeavyAtoms()
        denom = heavy if heavy else 1
        arom = sum(1 for atom in mol.GetAtoms() if atom.GetIsAromatic())
        hetero = sum(1 for atom in mol.GetAtoms() if atom.GetAtomicNum() not in (1, 6))
        halogen = sum(1 for atom in mol.GetAtoms() if atom.GetAtomicNum() in (9, 17, 35, 53))
        ring_count = int(mol.GetRingInfo().NumRings())
        rows.append(
            {
                "sample_index": int(idx),
                "SMILES": str(row[smiles_col]),
                "Tg_K": float(row[target_col]),
                "heavy_atoms": int(heavy),
                "ring_count": ring_count,
                "ring_containing": bool(ring_count > 0),
                "aromatic_fraction": float(arom / denom),
                "aromatic": bool(arom > 0),
                "heteroatom_fraction": float(hetero / denom),
                "halogen_fraction": float(halogen / denom),
            }
        )
        mols.append(mol)

    data = pd.DataFrame(rows).reset_index(drop=True)
    hetero_threshold = float(data["heteroatom_fraction"].median())
    data["heteroatom_rich"] = data["heteroatom_fraction"] >= hetero_threshold
    fps = [
        GetMorganFingerprintAsBitVect(mol, MORGAN_RADIUS, nBits=MORGAN_NBITS)
        for mol in mols
    ]
    return data, mols, fps, hetero_threshold


def cluster_by_cutoff(fps: list[DataStructs.cDataStructs.ExplicitBitVect], cutoff: float) -> list[list[int]]:
    clusters: list[list[int]] = []
    for i, fp in enumerate(fps):
        assigned = False
        for cluster in clusters:
            sim = DataStructs.TanimotoSimilarity(fp, fps[cluster[0]])
            if sim >= float(cutoff):
                cluster.append(i)
                assigned = True
                break
        if not assigned:
            clusters.append([i])
    return clusters


def make_cluster_table(data: pd.DataFrame, clusters: list[list[int]], cutoff: float) -> pd.DataFrame:
    rows = []
    for cluster_id, members in enumerate(clusters):
        subset = data.iloc[members]
        rows.append(
            {
                "cutoff": float(cutoff),
                "cluster_id": int(cluster_id),
                "cluster_size": int(len(members)),
                "member_indices": ";".join(str(int(i)) for i in members),
                "Tg_median_K": float(subset["Tg_K"].median()),
                "Tg_IQR_K": float(subset["Tg_K"].quantile(0.75) - subset["Tg_K"].quantile(0.25)),
                "ring_containing_fraction": float(subset["ring_containing"].mean()),
                "aromatic_fraction": float(subset["aromatic"].mean()),
                "heteroatom_fraction_median": float(subset["heteroatom_fraction"].median()),
                "representative_SMILES": str(subset.iloc[0]["SMILES"]),
            }
        )
    return pd.DataFrame(rows)


def cluster_splits_from_clusters(clusters: list[list[int]], seed: int) -> list[tuple[np.ndarray, np.ndarray]]:
    rng = np.random.RandomState(int(seed))
    shuffled = [list(cluster) for cluster in clusters]
    rng.shuffle(shuffled)
    fold_bins: list[list[int]] = [[] for _ in range(N_FOLDS)]
    fold_sizes = np.zeros(N_FOLDS, dtype=int)
    for cluster in sorted(shuffled, key=len, reverse=True):
        fold_idx = int(np.argmin(fold_sizes))
        fold_bins[fold_idx].extend(cluster)
        fold_sizes[fold_idx] += len(cluster)
    all_idx = np.arange(sum(len(cluster) for cluster in clusters))
    return [
        (np.setdiff1d(all_idx, np.array(sorted(fold), dtype=int)), np.array(sorted(fold), dtype=int))
        for fold in fold_bins
    ]


def max_tanimoto_test_to_train(
    fps: list[DataStructs.cDataStructs.ExplicitBitVect],
    train_idx: np.ndarray,
    test_idx: np.ndarray,
) -> np.ndarray:
    train_fps = [fps[i] for i in train_idx]
    out = np.zeros(len(test_idx), dtype=float)
    for j, i in enumerate(test_idx):
        sims = DataStructs.BulkTanimotoSimilarity(fps[int(i)], train_fps)
        out[j] = float(max(sims)) if sims else 0.0
    return out


def fold_and_test_records(
    data: pd.DataFrame,
    fps: list[DataStructs.cDataStructs.ExplicitBitVect],
    clusters_by_cutoff: dict[float, list[list[int]]],
    cluster_tables: dict[float, pd.DataFrame],
) -> tuple[pd.DataFrame, pd.DataFrame]:
    cluster_lookup = {}
    for cutoff, clusters in clusters_by_cutoff.items():
        for cluster_id, members in enumerate(clusters):
            for member in members:
                cluster_lookup[(cutoff, int(member))] = int(cluster_id)

    fold_rows = []
    test_rows = []
    for cutoff, clusters in clusters_by_cutoff.items():
        for seed in SEEDS:
            splits = cluster_splits_from_clusters(clusters, seed)
            for fold, (train_idx, test_idx) in enumerate(splits):
                subset = data.iloc[test_idx]
                smax = max_tanimoto_test_to_train(fps, train_idx, test_idx)
                fold_rows.append(
                    {
                        "cutoff": float(cutoff),
                        "seed": int(seed),
                        "fold": int(fold),
                        "fold_size": int(len(test_idx)),
                        "Tg_mean_K": float(subset["Tg_K"].mean()),
                        "Tg_median_K": float(subset["Tg_K"].median()),
                        "Tg_q25_K": float(subset["Tg_K"].quantile(0.25)),
                        "Tg_q75_K": float(subset["Tg_K"].quantile(0.75)),
                        "Tg_min_K": float(subset["Tg_K"].min()),
                        "Tg_max_K": float(subset["Tg_K"].max()),
                        "ring_containing_fraction": float(subset["ring_containing"].mean()),
                        "aromatic_fraction": float(subset["aromatic"].mean()),
                        "heteroatom_rich_fraction": float(subset["heteroatom_rich"].mean()),
                        "heteroatom_fraction_median": float(subset["heteroatom_fraction"].median()),
                        "halogen_fraction_median": float(subset["halogen_fraction"].median()),
                        "heavy_atoms_median": float(subset["heavy_atoms"].median()),
                        "Smax_median": float(np.median(smax)),
                        "Smax_q25": float(np.percentile(smax, 25)),
                        "Smax_q75": float(np.percentile(smax, 75)),
                        "low_Smax_tail_fraction": float(np.mean(smax < LOW_SMAX_THRESHOLD)),
                    }
                )
                for local_j, sample_idx in enumerate(test_idx):
                    row = data.iloc[int(sample_idx)]
                    test_rows.append(
                        {
                            "cutoff": float(cutoff),
                            "seed": int(seed),
                            "fold": int(fold),
                            "sample_index": int(row["sample_index"]),
                            "local_index": int(sample_idx),
                            "cluster_id": cluster_lookup[(cutoff, int(sample_idx))],
                            "SMILES": row["SMILES"],
                            "Tg_K": float(row["Tg_K"]),
                            "Smax": float(smax[local_j]),
                            "low_Smax_tail": bool(smax[local_j] < LOW_SMAX_THRESHOLD),
                            "ring_containing": bool(row["ring_containing"]),
                            "aromatic": bool(row["aromatic"]),
                            "heteroatom_fraction": float(row["heteroatom_fraction"]),
                            "heteroatom_rich": bool(row["heteroatom_rich"]),
                            "halogen_fraction": float(row["halogen_fraction"]),
                            "heavy_atoms": int(row["heavy_atoms"]),
                        }
                    )
    return pd.DataFrame(fold_rows), pd.DataFrame(test_rows)


def chemistry_fold_fraction_summary(folds: pd.DataFrame) -> str:
    parts = [
        ("Ring-containing", folds["ring_containing_fraction"] * 100),
        ("Aromatic", folds["aromatic_fraction"] * 100),
        ("Heteroatom-rich", folds["heteroatom_rich_fraction"] * 100),
    ]
    return "; ".join(f"{label} {fmt_iqr(values, decimals=1)}" for label, values in parts)


def manuscript_table(
    data: pd.DataFrame,
    cluster_tables: dict[float, pd.DataFrame],
    fold_df: pd.DataFrame,
    test_df: pd.DataFrame,
) -> pd.DataFrame:
    rows = []
    n = len(data)
    for cutoff in CUTOFFS:
        clusters = cluster_tables[cutoff]
        folds = fold_df[fold_df["cutoff"] == cutoff]
        tests = test_df[test_df["cutoff"] == cutoff]
        sizes = clusters["cluster_size"].to_numpy(int)
        singleton_clusters = int((sizes == 1).sum())
        singleton_units = int(sizes[sizes == 1].sum())
        low_tail = int(tests["low_Smax_tail"].sum())
        row = {
            "Cutoff": f"{cutoff:.2f}",
            "Clusters": int(len(clusters)),
            "Cluster size distribution": size_distribution_text(sizes),
            "Cluster size median [IQR]": fmt_iqr(sizes, decimals=1),
            "Largest cluster size": int(np.max(sizes)),
            "Singleton clusters n (%)": f"{singleton_clusters} ({pct(singleton_clusters, len(clusters)):.1f})",
            "Repeat units in singleton clusters n (%)": f"{singleton_units} ({pct(singleton_units, n):.1f})",
            "Fold size median [range]": f"{np.median(folds['fold_size']):.0f} [{int(folds['fold_size'].min())}, {int(folds['fold_size'].max())}]",
            "Fold Tg median range (K)": fmt_range(folds["Tg_median_K"], decimals=1),
            "Fold chemistry fraction median [IQR] (%)": chemistry_fold_fraction_summary(folds),
            "Smax median [IQR]": fmt_iqr(tests["Smax"], decimals=3),
            f"Low-Smax tail n (%) < {LOW_SMAX_THRESHOLD:.2f}": f"{low_tail} ({pct(low_tail, len(tests)):.1f})",
        }
        rows.append(row)
    return pd.DataFrame(rows)


def dataframe_to_markdown(df: pd.DataFrame) -> str:
    headers = [str(col) for col in df.columns]
    rows = [[str(value) for value in row] for row in df.to_numpy()]
    widths = [
        max(len(headers[i]), *(len(row[i]) for row in rows)) if rows else len(headers[i])
        for i in range(len(headers))
    ]
    header_line = "| " + " | ".join(headers[i].ljust(widths[i]) for i in range(len(headers))) + " |"
    separator = "| " + " | ".join("-" * widths[i] for i in range(len(headers))) + " |"
    body = [
        "| " + " | ".join(row[i].ljust(widths[i]) for i in range(len(headers)))
        + " |"
        for row in rows
    ]
    return "\n".join([header_line, separator, *body]) + "\n"


def latex_escape(value: object) -> str:
    text = str(value)
    for old, new in {
        "\\": r"\textbackslash{}",
        "&": r"\&",
        "%": r"\%",
        "$": r"\$",
        "#": r"\#",
        "_": r"\_",
        "{": r"\{",
        "}": r"\}",
        "~": r"\textasciitilde{}",
        "^": r"\textasciicircum{}",
    }.items():
        text = text.replace(old, new)
    return text


def dataframe_to_latex(df: pd.DataFrame, caption: str) -> str:
    cols = list(df.columns)
    lines = [
        r"\begin{table}[htbp]",
        r"\centering",
        f"\\caption{{{latex_escape(caption)}}}",
        r"\begin{tabular}{" + "l" * len(cols) + "}",
        r"\hline",
        " & ".join(latex_escape(col) for col in cols) + r" \\",
        r"\hline",
    ]
    for _, row in df.iterrows():
        lines.append(" & ".join(latex_escape(row[col]) for col in cols) + r" \\")
    lines.extend([r"\hline", r"\end{tabular}", r"\end{table}", ""])
    return "\n".join(lines)


def excel_col_name(index: int) -> str:
    name = ""
    while index:
        index, remainder = divmod(index - 1, 26)
        name = chr(65 + remainder) + name
    return name


def worksheet_xml(df: pd.DataFrame) -> str:
    rows = [list(df.columns)] + df.astype(str).values.tolist()
    xml_rows = []
    for r_idx, row in enumerate(rows, start=1):
        cells = []
        for c_idx, value in enumerate(row, start=1):
            ref = f"{excel_col_name(c_idx)}{r_idx}"
            cells.append(f'<c r="{ref}" t="inlineStr"><is><t>{escape(str(value))}</t></is></c>')
        xml_rows.append(f'<row r="{r_idx}">{"".join(cells)}</row>')
    return (
        '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
        '<worksheet xmlns="http://schemas.openxmlformats.org/spreadsheetml/2006/main">'
        f'<sheetData>{"".join(xml_rows)}</sheetData>'
        "</worksheet>"
    )


def write_xlsx(path: Path, sheets: dict[str, pd.DataFrame]) -> None:
    sheet_names = list(sheets.keys())
    with zipfile.ZipFile(path, "w", zipfile.ZIP_DEFLATED) as zf:
        zf.writestr(
            "[Content_Types].xml",
            '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
            '<Types xmlns="http://schemas.openxmlformats.org/package/2006/content-types">'
            '<Default Extension="rels" ContentType="application/vnd.openxmlformats-package.relationships+xml"/>'
            '<Default Extension="xml" ContentType="application/xml"/>'
            '<Override PartName="/xl/workbook.xml" ContentType="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet.main+xml"/>'
            + "".join(
                f'<Override PartName="/xl/worksheets/sheet{i}.xml" '
                'ContentType="application/vnd.openxmlformats-officedocument.spreadsheetml.worksheet+xml"/>'
                for i in range(1, len(sheet_names) + 1)
            )
            + "</Types>",
        )
        zf.writestr(
            "_rels/.rels",
            '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
            '<Relationships xmlns="http://schemas.openxmlformats.org/package/2006/relationships">'
            '<Relationship Id="rId1" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/officeDocument" Target="xl/workbook.xml"/>'
            "</Relationships>",
        )
        zf.writestr(
            "xl/workbook.xml",
            '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
            '<workbook xmlns="http://schemas.openxmlformats.org/spreadsheetml/2006/main" '
            'xmlns:r="http://schemas.openxmlformats.org/officeDocument/2006/relationships">'
            "<sheets>"
            + "".join(
                f'<sheet name="{escape(name[:31])}" sheetId="{i}" r:id="rId{i}"/>'
                for i, name in enumerate(sheet_names, start=1)
            )
            + "</sheets></workbook>",
        )
        zf.writestr(
            "xl/_rels/workbook.xml.rels",
            '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
            '<Relationships xmlns="http://schemas.openxmlformats.org/package/2006/relationships">'
            + "".join(
                f'<Relationship Id="rId{i}" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/worksheet" '
                f'Target="worksheets/sheet{i}.xml"/>'
                for i in range(1, len(sheet_names) + 1)
            )
            + "</Relationships>",
        )
        for i, name in enumerate(sheet_names, start=1):
            zf.writestr(f"xl/worksheets/sheet{i}.xml", worksheet_xml(sheets[name]))


def docx_text(text: str, style: str | None = None) -> str:
    style_xml = f'<w:pPr><w:pStyle w:val="{style}"/></w:pPr>' if style else ""
    return f"<w:p>{style_xml}<w:r><w:t>{escape(text)}</w:t></w:r></w:p>"


def write_minimal_docx(path: Path, manuscript: pd.DataFrame, caption: str) -> None:
    rows = [list(manuscript.columns)] + manuscript.astype(str).values.tolist()
    table_rows = []
    for row in rows:
        cells = "".join(
            "<w:tc><w:p><w:r><w:t>"
            + escape(str(value))
            + "</w:t></w:r></w:p></w:tc>"
            for value in row
        )
        table_rows.append(f"<w:tr>{cells}</w:tr>")
    document_xml = (
        '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
        '<w:document xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
        "<w:body>"
        + docx_text(TABLE_TITLE, "Heading1")
        + docx_text(caption)
        + f"<w:tbl>{''.join(table_rows)}</w:tbl>"
        + '<w:sectPr><w:pgSz w:w="15840" w:h="12240" w:orient="landscape"/></w:sectPr>'
        + "</w:body></w:document>"
    )
    with zipfile.ZipFile(path, "w", zipfile.ZIP_DEFLATED) as zf:
        zf.writestr(
            "[Content_Types].xml",
            '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
            '<Types xmlns="http://schemas.openxmlformats.org/package/2006/content-types">'
            '<Default Extension="rels" ContentType="application/vnd.openxmlformats-package.relationships+xml"/>'
            '<Default Extension="xml" ContentType="application/xml"/>'
            '<Override PartName="/word/document.xml" ContentType="application/vnd.openxmlformats-officedocument.wordprocessingml.document.main+xml"/>'
            "</Types>",
        )
        zf.writestr(
            "_rels/.rels",
            '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
            '<Relationships xmlns="http://schemas.openxmlformats.org/package/2006/relationships">'
            '<Relationship Id="rId1" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/officeDocument" Target="word/document.xml"/>'
            "</Relationships>",
        )
        zf.writestr("word/document.xml", document_xml)


def write_docx(path: Path, manuscript: pd.DataFrame, caption: str) -> None:
    try:
        from docx import Document
    except Exception:
        write_minimal_docx(path, manuscript, caption)
        return
    doc = Document()
    doc.add_heading(TABLE_TITLE, level=1)
    doc.add_paragraph(caption)
    table = doc.add_table(rows=1, cols=len(manuscript.columns))
    table.style = "Table Grid"
    columns = list(manuscript.columns)
    for i, col in enumerate(columns):
        table.rows[0].cells[i].text = col
    for _, row in manuscript.iterrows():
        cells = table.add_row().cells
        for i, col in enumerate(columns):
            cells[i].text = str(row[col])
    doc.save(path)


def boxplot_by_cutoff(ax, df: pd.DataFrame, col: str, ylabel: str, title: str) -> None:
    data = [df.loc[df["cutoff"] == cutoff, col].to_numpy(float) for cutoff in CUTOFFS]
    ax.boxplot(data, tick_labels=[f"{cutoff:.2f}" for cutoff in CUTOFFS], patch_artist=True)
    ax.set_xlabel("Morgan-Tanimoto cutoff")
    ax.set_ylabel(ylabel)
    ax.set_title(title)
    ax.grid(axis="y", alpha=0.25)


def plot_figure_s2(fold_df: pd.DataFrame, out_prefix: Path) -> list[Path]:
    plt.rcParams.update({"font.size": 9, "axes.titlesize": 10, "axes.labelsize": 9})
    fig, axes = plt.subplots(2, 2, figsize=(10.5, 7.5), constrained_layout=True)

    boxplot_by_cutoff(axes[0, 0], fold_df, "fold_size", "Fold test-set size", "A) Fold-size balance")
    boxplot_by_cutoff(axes[0, 1], fold_df, "Tg_median_K", "Fold median T$_g$ (K)", "B) Fold-level T$_g$ balance")
    boxplot_by_cutoff(axes[1, 0], fold_df, "Smax_median", "Fold median Smax", "C) Fold-level train-test similarity")

    ax = axes[1, 1]
    offsets = [-0.22, 0.0, 0.22]
    metrics = [
        ("ring_containing_fraction", "Ring-containing"),
        ("aromatic_fraction", "Aromatic"),
        ("heteroatom_rich_fraction", "Hetero-rich"),
    ]
    positions = np.arange(len(CUTOFFS)) + 1
    for offset, (col, label) in zip(offsets, metrics):
        data = [fold_df.loc[fold_df["cutoff"] == cutoff, col].to_numpy(float) * 100 for cutoff in CUTOFFS]
        bp = ax.boxplot(
            data,
            positions=positions + offset,
            widths=0.18,
            patch_artist=True,
            manage_ticks=False,
        )
        for patch in bp["boxes"]:
            patch.set_alpha(0.55)
        ax.plot([], [], label=label)
    ax.set_xticks(positions)
    ax.set_xticklabels([f"{cutoff:.2f}" for cutoff in CUTOFFS])
    ax.set_xlabel("Morgan-Tanimoto cutoff")
    ax.set_ylabel("Fold fraction (%)")
    ax.set_title("D) Fold-level chemistry balance")
    ax.grid(axis="y", alpha=0.25)
    ax.legend(frameon=False, fontsize=8)

    paths = []
    for ext, kwargs in {
        ".png": {"dpi": 300},
        ".svg": {},
        ".pdf": {},
        ".tiff": {"dpi": 600},
    }.items():
        path = out_prefix.with_suffix(ext)
        fig.savefig(path, **kwargs)
        paths.append(path)
    plt.close(fig)
    return paths


def main() -> None:
    OUT.mkdir(parents=True, exist_ok=True)
    FINAL.mkdir(parents=True, exist_ok=True)

    data, _mols, fps, hetero_threshold = load_dataset()
    clusters_by_cutoff = {cutoff: cluster_by_cutoff(fps, cutoff) for cutoff in CUTOFFS}
    cluster_tables = {
        cutoff: make_cluster_table(data, clusters, cutoff)
        for cutoff, clusters in clusters_by_cutoff.items()
    }
    fold_df, test_df = fold_and_test_records(data, fps, clusters_by_cutoff, cluster_tables)
    manuscript = manuscript_table(data, cluster_tables, fold_df, test_df)
    caption = table_s6_caption(hetero_threshold)
    cluster_groups = pd.concat(cluster_tables.values(), ignore_index=True)

    manuscript_csv = OUT / "TableS6_Cluster_Split_Characterization_manuscript.csv"
    manuscript_md = OUT / "TableS6_Cluster_Split_Characterization_manuscript.md"
    manuscript_tex = OUT / "TableS6_Cluster_Split_Characterization_manuscript.tex"
    xlsx_path = OUT / "TableS6_Cluster_Split_Characterization.xlsx"
    docx_path = OUT / "TableS6_Cluster_Split_Characterization.docx"
    cluster_groups_path = OUT / "TableS6_cluster_groups_raw.csv"
    fold_summary_path = OUT / "TableS6_fold_level_balance_raw.csv"
    test_annotations_path = OUT / "TableS6_test_record_smax_chemistry_annotations.csv"
    notes_path = OUT / "TableS6_FigureS2_cluster_split_notes.txt"
    figure_prefix = OUT / "FigureS2_cluster_split_fold_balance"

    manuscript.to_csv(manuscript_csv, index=False)
    manuscript_md.write_text(dataframe_to_markdown(manuscript), encoding="utf-8")
    manuscript_tex.write_text(dataframe_to_latex(manuscript, caption), encoding="utf-8")
    cluster_groups.to_csv(cluster_groups_path, index=False)
    fold_df.to_csv(fold_summary_path, index=False)
    test_df.to_csv(test_annotations_path, index=False)
    write_xlsx(
        xlsx_path,
        {
            "Table S6 manuscript": manuscript,
            "cluster groups": cluster_groups,
            "fold balance": fold_df,
            "test Smax chemistry": test_df,
        },
    )
    write_docx(docx_path, manuscript, caption)
    figure_paths = plot_figure_s2(fold_df, figure_prefix)

    for path in figure_paths:
        if path.suffix.lower() in {".png", ".tiff"}:
            shutil.copy2(path, FINAL / path.name)

    notes = f"""{TABLE_TITLE}
Generated: {datetime.now().isoformat()}

Inputs used:
- Processed benchmark data: {PROCESSED_CSV}
- Dataset metadata: {DATASET_META_JSON}
- Benchmark config: {CONFIG_YAML}
- Cluster split implementation: {SPLITS_PY}
- Smax implementation: {SIMILARITY_PY}

Method:
- Repeat-unit SMILES were parsed with RDKit.
- Morgan fingerprints were computed with radius={MORGAN_RADIUS}, nBits={MORGAN_NBITS}, matching configs/tg.yaml.
- Clusters were created with the benchmark's leader-style Tanimoto algorithm: each fingerprint is assigned to the first existing cluster whose representative fingerprint has similarity >= cutoff.
- Fold bins were generated exactly like src/cms_tg/splits.py for seeds {SEEDS}: shuffle clusters, sort by descending cluster size, then greedily assign the next cluster to the smallest fold.
- Smax is maximum Morgan-Tanimoto similarity from each test repeat unit to the training fold.
- Low-Smax tail fraction is defined as Smax < {LOW_SMAX_THRESHOLD:.2f}.
- Heteroatom-rich is defined using the dataset median heteroatom fraction: {hetero_threshold:.3f}.

Key values:
"""
    for _, row in manuscript.iterrows():
        notes += (
            f"- cutoff {row['Cutoff']}: {row['Clusters']} clusters; "
            f"{row['Singleton clusters n (%)']} singleton clusters; "
            f"fold size {row['Fold size median [range]']}; "
            f"Smax {row['Smax median [IQR]']}; "
            f"low-Smax tail {row[f'Low-Smax tail n (%) < {LOW_SMAX_THRESHOLD:.2f}']}.\n"
        )
    notes += """
Reviewer-facing interpretation:
The three cutoffs are not arbitrary in their effect: increasing the cutoff increases the number of clusters and singleton fraction, which increases structural stringency. Fold sizes remain tightly balanced because the benchmark greedily packs clusters by size into five folds after seed-specific shuffling. Tg and chemistry summaries in Figure S2 show whether that structural grouping preserves label and chemistry balance at the fold level.
"""
    notes_path.write_text(notes, encoding="utf-8")

    print(TABLE_TITLE)
    print(manuscript.to_string(index=False))
    print("Wrote:")
    for path in [
        manuscript_csv,
        manuscript_md,
        manuscript_tex,
        xlsx_path,
        docx_path,
        cluster_groups_path,
        fold_summary_path,
        test_annotations_path,
        notes_path,
        *figure_paths,
    ]:
        print(f"- {path}")


if __name__ == "__main__":
    main()
