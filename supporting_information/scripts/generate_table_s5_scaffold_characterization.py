from __future__ import annotations

from collections import Counter
from datetime import datetime
from html import escape
from pathlib import Path
import json
import zipfile

import numpy as np
import pandas as pd
from rdkit import Chem
from rdkit.Chem.Scaffolds import MurckoScaffold


import sys
from pathlib import Path

_SI_ROOT = Path(__file__).resolve().parents[1]
if str(_SI_ROOT) not in sys.path:
    sys.path.insert(0, str(_SI_ROOT))
from _paths import (  # noqa: E402
    BENCHMARK,
    CONFIG_YAML,
    DATASET_META_JSON,
    FINAL,
    INTERVALS_CSV,
    OUT,
    PROCESSED_CSV,
    RESULTS_CSV,
    ROOT,
    RUN,
    S8_MANIFEST,
    SIMILARITY_CSV,
    SUMMARY_CSV,
    XGB_SHAP_CSV,
    XGB_SHAP_FIG,
)


PROCESSED_CSV = RUN / "data" / "processed.csv"
DATASET_META_JSON = RUN / "data" / "dataset_meta.json"
SPLITS_PY = BENCHMARK / "src" / "cms_tg" / "splits.py"

TABLE_TITLE = "Table S5. Scaffold characterization for the polymer repeat-unit dataset"


def pct(numerator: float, denominator: float) -> float:
    return 100.0 * float(numerator) / float(denominator) if denominator else 0.0


def scaffold_for_mol(mol: Chem.Mol) -> str:
    return MurckoScaffold.MurckoScaffoldSmiles(mol=mol)


def generic_scaffold(scaffold_smiles: str) -> str:
    if not scaffold_smiles:
        return ""
    mol = Chem.MolFromSmiles(scaffold_smiles)
    if mol is None:
        return ""
    generic = MurckoScaffold.MakeScaffoldGeneric(mol)
    return Chem.MolToSmiles(generic, canonical=True)


def load_scaffold_annotations() -> pd.DataFrame:
    df = pd.read_csv(PROCESSED_CSV)
    with DATASET_META_JSON.open("r", encoding="utf-8") as handle:
        meta = json.load(handle)
    smiles_col = meta["smiles_col"]
    target_col = meta["target_col"]

    rows = []
    for idx, row in df.iterrows():
        smiles = str(row[smiles_col])
        mol = Chem.MolFromSmiles(smiles)
        if mol is None:
            rows.append(
                {
                    "sample_index": idx,
                    "SMILES": smiles,
                    "T_g (K)": row[target_col],
                    "valid_rdkit_mol": False,
                    "murcko_scaffold": "",
                    "generic_scaffold": "",
                    "is_empty_scaffold": True,
                    "repeat_unit_ring_count": np.nan,
                    "ring_containing_repeat_unit": False,
                    "scaffold_ring_count": np.nan,
                }
            )
            continue

        scaffold = scaffold_for_mol(mol)
        scaffold_mol = Chem.MolFromSmiles(scaffold) if scaffold else None
        rows.append(
            {
                "sample_index": idx,
                "SMILES": smiles,
                "T_g (K)": row[target_col],
                "valid_rdkit_mol": True,
                "murcko_scaffold": scaffold,
                "generic_scaffold": generic_scaffold(scaffold),
                "is_empty_scaffold": scaffold == "",
                "repeat_unit_ring_count": int(mol.GetRingInfo().NumRings()),
                "ring_containing_repeat_unit": bool(mol.GetRingInfo().NumRings() > 0),
                "scaffold_ring_count": (
                    int(scaffold_mol.GetRingInfo().NumRings()) if scaffold_mol is not None else 0
                ),
            }
        )
    return pd.DataFrame(rows)


def size_bin(size: int) -> str:
    if size == 1:
        return "1 (singleton)"
    if size == 2:
        return "2"
    if 3 <= size <= 5:
        return "3-5"
    if 6 <= size <= 10:
        return "6-10"
    return ">10"


def scaffold_group_table(annotations: pd.DataFrame) -> pd.DataFrame:
    grouped = (
        annotations.groupby("murcko_scaffold", dropna=False)
        .agg(
            n_repeat_units=("sample_index", "size"),
            generic_scaffold=("generic_scaffold", "first"),
            scaffold_ring_count=("scaffold_ring_count", "first"),
            representative_repeat_unit_smiles=("SMILES", "first"),
            tg_min_K=("T_g (K)", "min"),
            tg_median_K=("T_g (K)", "median"),
            tg_max_K=("T_g (K)", "max"),
        )
        .reset_index()
    )
    grouped["is_empty_scaffold"] = grouped["murcko_scaffold"].eq("")
    grouped["scaffold_size_bin"] = grouped["n_repeat_units"].map(size_bin)
    grouped = grouped.sort_values(
        ["n_repeat_units", "is_empty_scaffold", "murcko_scaffold"],
        ascending=[False, True, True],
    ).reset_index(drop=True)
    grouped.insert(0, "scaffold_rank_by_size", np.arange(1, len(grouped) + 1))
    return grouped


def summary_rows(annotations: pd.DataFrame, groups: pd.DataFrame) -> pd.DataFrame:
    n = len(annotations)
    valid = int(annotations["valid_rdkit_mol"].sum())
    empty_units = int(annotations["is_empty_scaffold"].sum())
    ring_units = int(annotations["ring_containing_repeat_unit"].sum())
    non_empty_groups = groups[~groups["is_empty_scaffold"]]
    singleton_scaffolds = int((groups["n_repeat_units"] == 1).sum())
    singleton_units = int(groups.loc[groups["n_repeat_units"] == 1, "n_repeat_units"].sum())
    generic_counts = groups.groupby("generic_scaffold", dropna=False)["n_repeat_units"].sum()
    n_generic_scaffolds = int(generic_counts.shape[0])
    non_empty_generic_counts = generic_counts[generic_counts.index != ""].sort_values(ascending=False)
    largest_generic = non_empty_generic_counts.index[0] if len(non_empty_generic_counts) else ""
    largest_generic_units = int(non_empty_generic_counts.iloc[0]) if len(non_empty_generic_counts) else 0
    empty_plus_largest_generic_units = empty_units + largest_generic_units

    rows = [
        {
            "Section": "Overall scaffold diagnostics",
            "Metric or scaffold": "Repeat units analyzed",
            "Value": f"{n}",
            "Fraction of dataset (%)": "100.0",
            "Definition / note": "Rows in the frozen processed benchmark dataset.",
        },
        {
            "Section": "Overall scaffold diagnostics",
            "Metric or scaffold": "Valid RDKit repeat-unit structures",
            "Value": f"{valid}",
            "Fraction of dataset (%)": f"{pct(valid, n):.1f}",
            "Definition / note": "Repeat-unit SMILES successfully parsed by RDKit.",
        },
        {
            "Section": "Overall scaffold diagnostics",
            "Metric or scaffold": "Unique Murcko scaffold groups, including empty scaffold",
            "Value": f"{len(groups)}",
            "Fraction of dataset (%)": "NA",
            "Definition / note": "Scaffold grouping exactly follows the scaffold split code.",
        },
        {
            "Section": "Overall scaffold diagnostics",
            "Metric or scaffold": "Unique non-empty Murcko scaffold groups",
            "Value": f"{len(non_empty_groups)}",
            "Fraction of dataset (%)": "NA",
            "Definition / note": "Excludes acyclic repeat units for which RDKit returns an empty Murcko scaffold.",
        },
        {
            "Section": "Overall scaffold diagnostics",
            "Metric or scaffold": "Unique generic Murcko scaffold groups",
            "Value": f"{n_generic_scaffolds}",
            "Fraction of dataset (%)": "NA",
            "Definition / note": "Scaffolds after RDKit MakeScaffoldGeneric atom/bond abstraction; empty scaffold retained as one group.",
        },
        {
            "Section": "Overall scaffold diagnostics",
            "Metric or scaffold": "Murcko-to-generic scaffold compression",
            "Value": f"{len(groups)} Murcko groups -> {n_generic_scaffolds} generic groups",
            "Fraction of dataset (%)": "NA",
            "Definition / note": "Generic abstraction collapses atom/bond identity and indicates how chemically specific the Murcko labels are.",
        },
        {
            "Section": "Overall scaffold diagnostics",
            "Metric or scaffold": "Singleton scaffold groups",
            "Value": f"{singleton_scaffolds}",
            "Fraction of dataset (%)": f"{pct(singleton_scaffolds, len(groups)):.1f}",
            "Definition / note": "Scaffold groups represented by one repeat unit; fraction denominator is unique scaffold groups.",
        },
        {
            "Section": "Overall scaffold diagnostics",
            "Metric or scaffold": "Repeat units assigned to singleton scaffolds",
            "Value": f"{singleton_units}",
            "Fraction of dataset (%)": f"{pct(singleton_units, n):.1f}",
            "Definition / note": "Dataset-level singleton scaffold fraction.",
        },
        {
            "Section": "Overall scaffold diagnostics",
            "Metric or scaffold": "Empty Murcko scaffold repeat units",
            "Value": f"{empty_units}",
            "Fraction of dataset (%)": f"{pct(empty_units, n):.1f}",
            "Definition / note": "Acyclic repeat units without a Murcko ring framework; the scaffold split groups these together.",
        },
        {
            "Section": "Overall scaffold diagnostics",
            "Metric or scaffold": "Largest non-empty generic scaffold class",
            "Value": f"{largest_generic or '[none]'}; {largest_generic_units} repeat units",
            "Fraction of dataset (%)": f"{pct(largest_generic_units, n):.1f}",
            "Definition / note": "Largest framework after generic Murcko abstraction; useful as a generic-framework concentration diagnostic.",
        },
        {
            "Section": "Overall scaffold diagnostics",
            "Metric or scaffold": "Empty scaffold plus largest generic scaffold class",
            "Value": f"{empty_plus_largest_generic_units} repeat units",
            "Fraction of dataset (%)": f"{pct(empty_plus_largest_generic_units, n):.1f}",
            "Definition / note": "Combined fraction of acyclic repeat units and the dominant generic ring framework class.",
        },
        {
            "Section": "Overall scaffold diagnostics",
            "Metric or scaffold": "Ring-containing repeat units",
            "Value": f"{ring_units}",
            "Fraction of dataset (%)": f"{pct(ring_units, n):.1f}",
            "Definition / note": "Repeat units with RDKit ring count > 0.",
        },
    ]
    return pd.DataFrame(rows)


def distribution_rows(groups: pd.DataFrame, n_units: int) -> pd.DataFrame:
    order = ["1 (singleton)", "2", "3-5", "6-10", ">10"]
    rows = []
    for label in order:
        subset = groups[groups["scaffold_size_bin"] == label]
        n_scaffolds = int(len(subset))
        units = int(subset["n_repeat_units"].sum())
        rows.append(
            {
                "Section": "Scaffold-size distribution",
                "Metric or scaffold": label,
                "Value": f"{n_scaffolds} scaffold groups; {units} repeat units",
                "Fraction of dataset (%)": f"{pct(units, n_units):.1f}",
                "Definition / note": f"Scaffold group size bin; {pct(n_scaffolds, len(groups)):.1f}% of scaffold groups.",
            }
        )
    return pd.DataFrame(rows)


def representative_rows(groups: pd.DataFrame, n_units: int, n_examples: int = 10) -> pd.DataFrame:
    top = groups.sort_values(
        ["n_repeat_units", "is_empty_scaffold", "murcko_scaffold"],
        ascending=[False, True, True],
    ).head(n_examples)
    rows = []
    for _, row in top.iterrows():
        label = row["murcko_scaffold"] if row["murcko_scaffold"] else "[empty scaffold]"
        rows.append(
            {
                "Section": "Representative scaffold examples",
                "Metric or scaffold": label,
                "Value": f"{int(row['n_repeat_units'])} repeat units",
                "Fraction of dataset (%)": f"{pct(row['n_repeat_units'], n_units):.1f}",
                "Definition / note": (
                    f"Generic scaffold: {row['generic_scaffold'] or '[empty]'}; "
                    f"example repeat unit: {row['representative_repeat_unit_smiles']}; "
                    f"Tg range: {row['tg_min_K']:.1f}-{row['tg_max_K']:.1f} K."
                ),
            }
        )
    return pd.DataFrame(rows)


def manuscript_table(annotations: pd.DataFrame, groups: pd.DataFrame) -> pd.DataFrame:
    n_units = len(annotations)
    return pd.concat(
        [
            summary_rows(annotations, groups),
            distribution_rows(groups, n_units),
            representative_rows(groups, n_units),
        ],
        ignore_index=True,
    )


def dataframe_to_markdown(df: pd.DataFrame) -> str:
    headers = [str(col) for col in df.columns]
    rows = [[str(value) for value in row] for row in df.to_numpy()]
    widths = [
        max(len(headers[i]), *(len(row[i]) for row in rows)) if rows else len(headers[i])
        for i in range(len(headers))
    ]
    header_line = "| " + " | ".join(headers[i].ljust(widths[i]) for i in range(len(headers))) + " |"
    separator = "| " + " | ".join("-" * widths[i] for i in range(len(headers))) + " |"
    body = [
        "| " + " | ".join(row[i].ljust(widths[i]) for i in range(len(headers))) + " |"
        for row in rows
    ]
    return "\n".join([header_line, separator, *body]) + "\n"


def latex_escape(value: object) -> str:
    text = str(value)
    for old, new in {
        "\\": r"\textbackslash{}",
        "&": r"\&",
        "%": r"\%",
        "$": r"\$",
        "#": r"\#",
        "_": r"\_",
        "{": r"\{",
        "}": r"\}",
        "~": r"\textasciitilde{}",
        "^": r"\textasciicircum{}",
    }.items():
        text = text.replace(old, new)
    return text


def dataframe_to_latex(df: pd.DataFrame) -> str:
    cols = list(df.columns)
    lines = [
        r"\begin{table}[htbp]",
        r"\centering",
        f"\\caption{{{latex_escape(TABLE_TITLE)}}}",
        r"\begin{tabular}{" + "l" * len(cols) + "}",
        r"\hline",
        " & ".join(latex_escape(col) for col in cols) + r" \\",
        r"\hline",
    ]
    for _, row in df.iterrows():
        lines.append(" & ".join(latex_escape(row[col]) for col in cols) + r" \\")
    lines.extend([r"\hline", r"\end{tabular}", r"\end{table}", ""])
    return "\n".join(lines)


def excel_col_name(index: int) -> str:
    name = ""
    while index:
        index, remainder = divmod(index - 1, 26)
        name = chr(65 + remainder) + name
    return name


def worksheet_xml(df: pd.DataFrame) -> str:
    rows = [list(df.columns)] + df.astype(str).values.tolist()
    xml_rows = []
    for r_idx, row in enumerate(rows, start=1):
        cells = []
        for c_idx, value in enumerate(row, start=1):
            ref = f"{excel_col_name(c_idx)}{r_idx}"
            cells.append(f'<c r="{ref}" t="inlineStr"><is><t>{escape(str(value))}</t></is></c>')
        xml_rows.append(f'<row r="{r_idx}">{"".join(cells)}</row>')
    return (
        '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
        '<worksheet xmlns="http://schemas.openxmlformats.org/spreadsheetml/2006/main">'
        f'<sheetData>{"".join(xml_rows)}</sheetData>'
        "</worksheet>"
    )


def write_xlsx(path: Path, sheets: dict[str, pd.DataFrame]) -> None:
    sheet_names = list(sheets.keys())
    with zipfile.ZipFile(path, "w", zipfile.ZIP_DEFLATED) as zf:
        zf.writestr(
            "[Content_Types].xml",
            '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
            '<Types xmlns="http://schemas.openxmlformats.org/package/2006/content-types">'
            '<Default Extension="rels" ContentType="application/vnd.openxmlformats-package.relationships+xml"/>'
            '<Default Extension="xml" ContentType="application/xml"/>'
            '<Override PartName="/xl/workbook.xml" ContentType="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet.main+xml"/>'
            + "".join(
                f'<Override PartName="/xl/worksheets/sheet{i}.xml" '
                'ContentType="application/vnd.openxmlformats-officedocument.spreadsheetml.worksheet+xml"/>'
                for i in range(1, len(sheet_names) + 1)
            )
            + "</Types>",
        )
        zf.writestr(
            "_rels/.rels",
            '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
            '<Relationships xmlns="http://schemas.openxmlformats.org/package/2006/relationships">'
            '<Relationship Id="rId1" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/officeDocument" Target="xl/workbook.xml"/>'
            "</Relationships>",
        )
        zf.writestr(
            "xl/workbook.xml",
            '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
            '<workbook xmlns="http://schemas.openxmlformats.org/spreadsheetml/2006/main" '
            'xmlns:r="http://schemas.openxmlformats.org/officeDocument/2006/relationships">'
            "<sheets>"
            + "".join(
                f'<sheet name="{escape(name[:31])}" sheetId="{i}" r:id="rId{i}"/>'
                for i, name in enumerate(sheet_names, start=1)
            )
            + "</sheets></workbook>",
        )
        zf.writestr(
            "xl/_rels/workbook.xml.rels",
            '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
            '<Relationships xmlns="http://schemas.openxmlformats.org/package/2006/relationships">'
            + "".join(
                f'<Relationship Id="rId{i}" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/worksheet" '
                f'Target="worksheets/sheet{i}.xml"/>'
                for i in range(1, len(sheet_names) + 1)
            )
            + "</Relationships>",
        )
        for i, name in enumerate(sheet_names, start=1):
            zf.writestr(f"xl/worksheets/sheet{i}.xml", worksheet_xml(sheets[name]))


def docx_text(text: str, style: str | None = None) -> str:
    style_xml = f'<w:pPr><w:pStyle w:val="{style}"/></w:pPr>' if style else ""
    return f"<w:p>{style_xml}<w:r><w:t>{escape(text)}</w:t></w:r></w:p>"


def write_minimal_docx(path: Path, manuscript: pd.DataFrame) -> None:
    rows = [list(manuscript.columns)] + manuscript.astype(str).values.tolist()
    table_rows = []
    for row in rows:
        cells = "".join(
            "<w:tc><w:p><w:r><w:t>"
            + escape(str(value))
            + "</w:t></w:r></w:p></w:tc>"
            for value in row
        )
        table_rows.append(f"<w:tr>{cells}</w:tr>")
    document_xml = (
        '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
        '<w:document xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
        "<w:body>"
        + docx_text(TABLE_TITLE, "Heading1")
        + docx_text(
            "Murcko scaffolds were computed with the same RDKit call used by the scaffold split. "
            "Empty scaffolds correspond to acyclic repeat units and are intentionally reported because "
            "they form a single group in the scaffold-split diagnostic."
        )
        + f"<w:tbl>{''.join(table_rows)}</w:tbl>"
        + '<w:sectPr><w:pgSz w:w="15840" w:h="12240" w:orient="landscape"/></w:sectPr>'
        + "</w:body></w:document>"
    )
    with zipfile.ZipFile(path, "w", zipfile.ZIP_DEFLATED) as zf:
        zf.writestr(
            "[Content_Types].xml",
            '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
            '<Types xmlns="http://schemas.openxmlformats.org/package/2006/content-types">'
            '<Default Extension="rels" ContentType="application/vnd.openxmlformats-package.relationships+xml"/>'
            '<Default Extension="xml" ContentType="application/xml"/>'
            '<Override PartName="/word/document.xml" ContentType="application/vnd.openxmlformats-officedocument.wordprocessingml.document.main+xml"/>'
            "</Types>",
        )
        zf.writestr(
            "_rels/.rels",
            '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
            '<Relationships xmlns="http://schemas.openxmlformats.org/package/2006/relationships">'
            '<Relationship Id="rId1" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/officeDocument" Target="word/document.xml"/>'
            "</Relationships>",
        )
        zf.writestr("word/document.xml", document_xml)


def write_docx(path: Path, manuscript: pd.DataFrame) -> None:
    try:
        from docx import Document
    except Exception:
        write_minimal_docx(path, manuscript)
        return

    doc = Document()
    doc.add_heading(TABLE_TITLE, level=1)
    doc.add_paragraph(
        "Murcko scaffolds were computed with the same RDKit call used by the scaffold split. "
        "Empty scaffolds correspond to acyclic repeat units and are intentionally reported because "
        "they form a single group in the scaffold-split diagnostic."
    )
    table = doc.add_table(rows=1, cols=len(manuscript.columns))
    table.style = "Table Grid"
    columns = list(manuscript.columns)
    for i, col in enumerate(columns):
        table.rows[0].cells[i].text = col
    for _, row in manuscript.iterrows():
        cells = table.add_row().cells
        for i, col in enumerate(columns):
            cells[i].text = str(row[col])
    doc.save(path)


def main() -> None:
    OUT.mkdir(parents=True, exist_ok=True)
    annotations = load_scaffold_annotations()
    groups = scaffold_group_table(annotations)
    manuscript = manuscript_table(annotations, groups)

    annotations_path = OUT / "TableS5_scaffold_annotations_per_repeat_unit.csv"
    groups_path = OUT / "TableS5_scaffold_groups_raw.csv"
    manuscript_csv = OUT / "TableS5_Scaffold_Characterization_manuscript.csv"
    manuscript_md = OUT / "TableS5_Scaffold_Characterization_manuscript.md"
    manuscript_tex = OUT / "TableS5_Scaffold_Characterization_manuscript.tex"
    xlsx_path = OUT / "TableS5_Scaffold_Characterization.xlsx"
    docx_path = OUT / "TableS5_Scaffold_Characterization.docx"
    notes_path = OUT / "TableS5_scaffold_characterization_notes.txt"

    annotations.to_csv(annotations_path, index=False)
    groups.to_csv(groups_path, index=False)
    manuscript.to_csv(manuscript_csv, index=False)
    manuscript_md.write_text(dataframe_to_markdown(manuscript), encoding="utf-8")
    manuscript_tex.write_text(dataframe_to_latex(manuscript), encoding="utf-8")
    write_xlsx(
        xlsx_path,
        {
            "Table S5 manuscript": manuscript,
            "scaffold groups": groups,
            "per repeat unit": annotations,
        },
    )
    write_docx(docx_path, manuscript)

    n_units = len(annotations)
    empty_units = int(annotations["is_empty_scaffold"].sum())
    ring_units = int(annotations["ring_containing_repeat_unit"].sum())
    singleton_scaffolds = int((groups["n_repeat_units"] == 1).sum())
    singleton_units = int(groups.loc[groups["n_repeat_units"] == 1, "n_repeat_units"].sum())
    top_group = groups.iloc[0]
    generic_counts = groups.groupby("generic_scaffold", dropna=False)["n_repeat_units"].sum()
    non_empty_generic_counts = generic_counts[generic_counts.index != ""].sort_values(ascending=False)
    largest_generic = non_empty_generic_counts.index[0] if len(non_empty_generic_counts) else ""
    largest_generic_units = int(non_empty_generic_counts.iloc[0]) if len(non_empty_generic_counts) else 0

    notes = f"""Table S5 scaffold characterization
Generated: {datetime.now().isoformat()}

Inputs used:
- Processed benchmark data: {PROCESSED_CSV}
- Dataset metadata: {DATASET_META_JSON}
- Scaffold split implementation: {SPLITS_PY}

Method:
- Repeat-unit SMILES were parsed with RDKit.
- Murcko scaffolds were computed using MurckoScaffold.MurckoScaffoldSmiles(mol=m), matching src/cms_tg/splits.py.
- Empty Murcko scaffold strings are retained, because the scaffold split groups acyclic repeat units under the same empty scaffold key.
- Generic scaffolds were computed with MurckoScaffold.MakeScaffoldGeneric for characterization only; they are not used to make the benchmark splits.

Key values:
- Repeat units analyzed: {n_units}
- Unique Murcko scaffold groups including empty scaffold: {len(groups)}
- Unique non-empty Murcko scaffold groups: {int((~groups['is_empty_scaffold']).sum())}
- Singleton scaffold groups: {singleton_scaffolds} ({pct(singleton_scaffolds, len(groups)):.1f}% of scaffold groups)
- Repeat units assigned to singleton scaffolds: {singleton_units} ({pct(singleton_units, n_units):.1f}% of dataset)
- Empty Murcko scaffold repeat units: {empty_units} ({pct(empty_units, n_units):.1f}% of dataset)
- Largest non-empty generic scaffold class: {largest_generic or '[none]'} with {largest_generic_units} repeat units ({pct(largest_generic_units, n_units):.1f}% of dataset)
- Empty scaffold plus largest generic scaffold class: {empty_units + largest_generic_units} repeat units ({pct(empty_units + largest_generic_units, n_units):.1f}% of dataset)
- Ring-containing repeat units: {ring_units} ({pct(ring_units, n_units):.1f}% of dataset)
- Largest scaffold group: {top_group['murcko_scaffold'] or '[empty scaffold]'} with {int(top_group['n_repeat_units'])} repeat units

Reviewer-facing interpretation:
This table supports describing the scaffold split as a framework-leakage diagnostic rather than a full polymer out-of-distribution benchmark. A non-trivial empty-scaffold/acyclic fraction and many singleton scaffold groups mean the split stresses repeat-unit framework memorization, but it does not fully encode polymer architecture, sequence, molecular weight distribution, processing history, or higher-order morphology.
"""
    notes_path.write_text(notes, encoding="utf-8")

    print(TABLE_TITLE)
    print(f"Repeat units analyzed: {n_units}")
    print(f"Unique scaffold groups including empty: {len(groups)}")
    print(f"Unique non-empty scaffold groups: {int((~groups['is_empty_scaffold']).sum())}")
    print(f"Singleton scaffold groups: {singleton_scaffolds} ({pct(singleton_scaffolds, len(groups)):.1f}% of scaffold groups)")
    print(f"Repeat units in singleton scaffolds: {singleton_units} ({pct(singleton_units, n_units):.1f}% of dataset)")
    print(f"Empty Murcko scaffold repeat units: {empty_units} ({pct(empty_units, n_units):.1f}% of dataset)")
    print(f"Largest non-empty generic scaffold class: {largest_generic or '[none]'} ({largest_generic_units} repeat units; {pct(largest_generic_units, n_units):.1f}% of dataset)")
    print(f"Empty + largest generic scaffold class: {empty_units + largest_generic_units} repeat units ({pct(empty_units + largest_generic_units, n_units):.1f}% of dataset)")
    print(f"Ring-containing repeat units: {ring_units} ({pct(ring_units, n_units):.1f}% of dataset)")
    print("Wrote:")
    for path in [
        manuscript_csv,
        manuscript_md,
        manuscript_tex,
        xlsx_path,
        docx_path,
        groups_path,
        annotations_path,
        notes_path,
    ]:
        print(f"- {path}")


if __name__ == "__main__":
    main()
